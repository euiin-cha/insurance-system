# ============================================================
# 한화손해보험 상품 비교분석 시스템  |  app.py
# ============================================================
import streamlit as st
import os, io, re, json, copy
import pdfplumber
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from collections import defaultdict, Counter
from difflib import SequenceMatcher

# ── 페이지 설정 ──────────────────────────────────────────────
st.set_page_config(
    page_title="한화손보 AI기반 상품 비교분석시스템",
    page_icon="🏢", layout="wide"
)

st.markdown("""
<style>
/* ── 사이드바 전체 주황 배경 ── */
section[data-testid="stSidebar"]{min-width:220px!important;max-width:220px!important}
section[data-testid="stSidebar"] > div:first-child{
    background:#FF6B00!important;padding:0!important;margin:0!important;
    border:none!important;border-radius:0!important;}
section[data-testid="stSidebar"] *{color:#fff!important;}
section[data-testid="stSidebar"] hr{border-color:rgba(255,255,255,.3)!important;}
section[data-testid="stSidebar"] .stButton>button{
    text-align:left!important;justify-content:flex-start!important;
    border-radius:8px!important;font-size:13px!important;
    padding:8px 12px!important;width:100%!important;margin-bottom:3px!important;
    border:1.5px solid rgba(255,255,255,.4)!important;
    background:rgba(255,255,255,.15)!important;color:#fff!important;}
section[data-testid="stSidebar"] .stButton>button[kind="primary"]{
    background:#fff!important;color:#FF6B00!important;
    border:1.5px solid #fff!important;font-weight:700!important;}
section[data-testid="stSidebar"] .stButton>button:hover{
    background:rgba(255,255,255,.3)!important;color:#fff!important;}
/* ── 상단 헤더바 흰색 ── */
header[data-testid="stHeader"]{background:#ffffff!important;}
header[data-testid="stHeader"] *{color:#333!important;}
div[data-testid="stToolbar"]{background:#ffffff!important;}
/* ── 메인 콘텐츠 ── */
.main .block-container{
    padding-top:1.2rem!important;
    padding-left:1.5rem;padding-right:1.5rem;}
/* ── Streamlit 기본 요소 ── */
.stTabs [data-baseweb="tab"]{font-size:15px;font-weight:600;padding:8px 20px;}
.stTabs [aria-selected="true"]{
    color:#FF6B00!important;border-bottom:3px solid #FF6B00!important;font-weight:700!important;}
.stTabs [data-baseweb="tab"]:hover{color:#FF6B00!important;}
.badge{display:inline-block;padding:2px 8px;border-radius:12px;font-size:12px;font-weight:600;margin:1px;}
.upload-card{border:1px solid #eee;border-radius:8px;padding:14px 16px;
    background:#fff;margin-bottom:8px;}
.upload-card-active{border-color:#FF6B00;background:#FFF8F3;}
</style>
""", unsafe_allow_html=True)

# ════════════════════════════════════════════════════════════
#  A. 엑셀 생성 로직 — 검증된 함수들
# ════════════════════════════════════════════════════════════

CHEZUNG_PAT = re.compile(r'\s*\(체증형[^)]*\)', re.IGNORECASE)

KNOWN_COMPANIES = ["한화손해보험","삼성화재","현대해상","메리츠화재",
                   "KB손해보험","DB손해보험","흥국화재","롯데손해보험"]

# ── A-1: 담보명 조합 ─────────────────────────────────────────
def build_name(raw, company):
    lines = [l.strip() for l in str(raw).split('\n') if l.strip()]
    if not lines: return ''
    if company in ("한화손해보험", "현대해상"):
        result = lines[0]; seen = {lines[0]}
        for line in lines[1:]:
            if line in seen: break
            seen.add(line)
            result += line if (line.startswith('(') or result.endswith('(')) else line
        return result.strip()
    result = lines[0]
    for line in lines[1:]:
        if line.startswith('(') or result.endswith('('): result += line
        else: result += line
    return result.strip()

# ── A-2: 담보명 정제 ─────────────────────────────────────────
def clean_name(name):
    name = re.sub(r'\[갱신형\]|\[비갱신형\]', '', name)
    m = re.match(r'^(.+?)\(간편\)\1.*$', name)
    if m: return m.group(1) + '(간편)'
    m2 = re.match(r'^(.*?\(간편\))[가-힣\[]', name)
    if m2 and len(m2.group(1)) < len(name) * 0.7: return m2.group(1)
    for n in range(max(5, len(name)//3), len(name)//2 + 1):
        prefix = name[:n]
        if len(prefix) >= 5 and prefix in name[n:]:
            return name[:n].strip()
    name = re.sub(r'\s*\(비갱신형/갱신형\)', '', name)
    name = re.sub(r'\s*\(비갱신형\)', '', name)
    name = re.sub(r'\s*\(갱신형\)', '', name)
    name = re.sub(r'\s*\(갱신형포함\)', '', name)
    return name.strip()

# ── A-3: 유효성 검증 ────────────────────────────────────────
def is_valid(name):
    if not name or len(name) < 5: return False
    if not re.search(r'[가-힣]', name): return False
    if name.count('(') != name.count(')'): return False
    if re.match(r'^\(', name): return False
    if re.match(r'^\d', name): return False
    if re.search(r'\d$', name): return False
    if name.endswith(','): return False
    if re.search(r'합니다\.?$|합니다\s*$', name): return False
    if re.search(r'보장되지 않습니다|지급되지 않습니다', name): return False
    if '※' in name: return False
    if re.match(r'^[①②③④⑤⑥⑦⑧⑨]', name): return False
    # 조사/어미로 끝나는 짧은 잘린 담보명
    if len(name) < 10 and re.search(r'(의|에|이|하|받|진|을|를|은|는|과|와|로|서|도|만)$', name):
        return False
    return True

# ── A-4: 비갱신형/갱신형 판단 ───────────────────────────────
def get_renewal_flags(name_orig):
    has_non = bool(re.search(r'비갱신형', name_orig))
    has_ren = bool(re.search(r'갱신형', name_orig))
    if (has_non and has_ren) or '갱신형포함' in name_orig:
        return "V", "V"
    elif has_ren and not has_non:
        return "", "V"
    else:
        return "V", ""

# ── A-5: 지급사유 정제 ──────────────────────────────────────
def clean_reason(reason):
    reason = re.sub(r'\s*\(\s*90일\s*미만.*', '', reason, flags=re.DOTALL)
    reason = re.sub(r'\s*\(\s*1년\s*미만.*', '', reason, flags=re.DOTALL)
    reason = re.sub(r',\s*90일\s*이상.*지급[^)]*', '', reason, flags=re.DOTALL)
    reason = re.sub(r'(단[,，]?\s*)?(계약일|가입후)[^。\.]*?(미만|이상)[^。\.]*?%.*', '', reason)
    reason = re.sub(r'(보험가입금액|가입금액)[^\s,。\.]*', '', reason)
    return re.sub(r'\s+', ' ', reason).strip()

def is_valid_reason(reason):
    if not reason: return False
    if re.match(r'^(가입금액|보험가입금액|특약가입금액|\d)', reason): return False
    if re.search(r'합니다$|보장되지 않습니다|다음날로 합니다', reason): return False
    return any(kw in reason for kw in
        ['경우','때','받은','발생','확정','사망','입원','수술','진단','해당','이용'])

# ── A-6: 감액/면책 추출 ─────────────────────────────────────
def extract_reduction(name, reason_raw):
    combined = name + ' ' + reason_raw
    if re.search(r'1년50%|1년감액|1년미만50%', name): return "1년 미만 50%"
    if re.search(r'90일50%|90일감액', name): return "90일 미만 50%"
    if re.search(r'90일\s*미만.*10%', combined): return "90일 미만 10%, 1년 미만 50%"
    if re.search(r'1년\s*미만.*50%', combined): return "1년 미만 50%"
    return None

def extract_exemption(name, reason_raw):
    combined = name + ' ' + reason_raw
    if re.search(r'90일면책', name): return "90일"
    if re.search(r'1년면책', name): return "1년"
    if re.search(r'90일.*다음날|보장개시.*90일', combined): return "90일"
    if re.search(r'1년.*면책|1년이 지난.*다음날', combined): return "1년"
    return None

# ── A-7: 지급횟수 판단 ──────────────────────────────────────
def guess_freq(name, reason_raw=""):
    n = name + ' ' + reason_raw
    if '진단비' in name:
        if '연간1회한' in n or '연간 1회한' in n: return "연간 1회한"
        if '5회한' in n: return "최초 5회한"
        return "최초 1회"
    if '수술비' in name: return "수술 1회당"
    if '입원일당' in name or '입원비' in name:
        m = re.search(r'(\d+)일\s*한도', n)
        return f"1일당({m.group(1)}일 한도)" if m else "1일당"
    if '1사고당' in n: return "1사고당"
    if '연간1회한' in n or '연간 1회한' in n: return "연간 1회한"
    if '5회한' in n: return "최초 5회한"
    if '통원' in name: return "1회당"
    return "최초 1회"

# ── A-8: 분류 ────────────────────────────────────────────────
def categorize(name):
    if any(k in name for k in ['암진단','암수술','유사암','재진단암','전이암','항암방사선',
        '항암약물','항암중입자','암직접치료','암요양병원','암통합','암입원','암주요치료',
        '암치료비','표적항암','면역항암','암(유사암','암(4대유사','하이클래스암',
        '계속받는항암','암주요검사','갑상선암진단']): return '암'
    if any(k in name for k in ['뇌졸중','뇌출혈','뇌혈관','뇌경색','급성심근경색',
        '허혈성심장','혈전용해','혈전제거','순환계','특정상해성뇌출혈','심혈관',
        '심뇌혈관','뇌심장','외상성뇌']): return '뇌/심'
    if any(k in name for k in ['수술비','1~5종수술','1-5종수술','골절수술','화상수술',
        '각막이식수술','다빈치로봇','흉터복원수술','심장수술비','다발성질병수술',
        '5대기관질병수술','심뇌혈관질환수술']): return '수술비'
    if any(k in name for k in ['입원일당','입원비','중환자실입원','요양병원입원']):
        return '입원비'
    if any(k in name for k in ['통합치료비','치료비','혈전용해치료','간병인사용',
        '간호·간병통합','재활치료','방사선치료','약물치료','산정특례','주요치료비',
        '심뇌혈관질환주요약제','하이클래스암특정치료']): return '치료비'
    if any(k in name for k in ['통원','외래','응급실']): return '통원비'
    return '기타'

# ── A-9: 지급사유 추론 ──────────────────────────────────────
def guess_reason(name):
    if '사망' in name: return "보험기간 중 해당 사유로 사망한 경우"
    if '암진단' in name or ('암' in name and '진단' in name and '재진단' not in name):
        return "보험기간 중 암 진단확정 시(최초 1회)"
    if '유사암' in name and '진단' in name: return "보험기간 중 유사암 진단확정 시(각 최초 1회)"
    if '재진단암' in name: return "재진단암보장개시일 이후 재진단암 진단확정 시(5회 한도)"
    if '뇌졸중' in name and '진단' in name: return "보험기간 중 뇌졸중 진단확정 시(최초 1회)"
    if '급성심근경색' in name: return "보험기간 중 급성심근경색증 진단확정 시(최초 1회)"
    if '후유장해' in name: return "보험기간 중 후유장해 발생 시"
    if '수술비' in name: return "보험기간 중 해당 사유로 수술 받은 경우"
    if '입원일당' in name or '입원비' in name:
        return "보험기간 중 해당 사유로 입원하여 치료받은 경우"
    if '항암' in name: return "보험기간 중 암 진단 후 항암치료 받은 경우"
    if '치매' in name: return "보험기간 중 치매 진단확정 시(최초 1회)"
    if '납입면제' in name: return "납입면제 사유 발생 시(최초 1회)"
    if '통원' in name: return "보험기간 중 해당 사유로 통원하여 치료받은 경우"
    if '치료비' in name: return "보험기간 중 해당 치료를 받은 경우"
    if '진단비' in name: return "보험기간 중 해당 질환 진단확정 시"
    if '간병인' in name: return "보험기간 중 입원 중 간병인 사용 시"
    return ""

# ── A-10: 상품구조 추출 ─────────────────────────────────────
def extract_product_structure(pdf_text, company):
    """PDF 원문에서 상품구조 정보 정확 추출"""
    result = {
        "company": company,
        "product_name": "미기재", "jong": "미기재", "hyung": "미기재",
        "basic_policy": "미기재", "insurance_period": "미기재",
        "payment_period": "미기재", "age_range": "미기재",
        "premium_waiver": "미기재", "no_claim": "미기재",
        "services": "해당없음", "discounts": "해당없음",
    }

    # ── 상품명: 첫 페이지 두 번째 줄에서 "상품요약서" 제거 ────
    first_lines = [l.strip() for l in pdf_text[:2000].split('\n')
                   if l.strip() and len(l.strip()) > 5]
    for line in first_lines[:10]:
        if re.search(r'건강보험|손해보험|생명보험|화재보험', line) \
           and "특별약관" not in line and "적용" not in line:
            # "XXX 상품요약서" → "XXX" 로 정제
            name = re.sub(r'\s*상품요약서\s*$', '', line).strip()
            name = re.sub(r'\s*요약서\s*$', '', name).strip()
            result["product_name"] = name[:100]
            break

    # ── 종(種): "보험종목의 세목" 섹션 우선, 없으면 전체에서 추출 ─
    jong_section = ""
    for kw in ["보험종목의 세목", "가입자격"]:
        ki = pdf_text.find(kw)
        if ki >= 0:
            jong_section = pdf_text[ki:ki+4000]; break

    # 세목 섹션이 없거나 짧으면 전체에서 탐색 (삼성화재 등)
    jong_matches_all = re.findall(r'([1-9]종\([^)]{5,80}\))', jong_section or pdf_text)
    seen_jong_num, jong_list = set(), []
    for jm in jong_matches_all:
        num = re.match(r'([1-9])종', jm).group(1)
        if num not in seen_jong_num:
            seen_jong_num.add(num)
            jong_list.append(jm)
    # 세목 섹션 패턴 미발견 시 단순 "N종" 라인에서 보완
    if not jong_list:
        # "A1) 1종(설명)" 또는 "1종(설명)" 형태
        jong_simple = re.findall(r'(?:A\d\)\s*)?([1-9]종\([^)]{3,80}\))', pdf_text)
        for jm in jong_simple:
            num = re.match(r'([1-9])종', jm).group(1)
            if num not in seen_jong_num:
                seen_jong_num.add(num)
                jong_list.append(jm)
    result["jong"] = " / ".join(jong_list) if jong_list else "미기재"

    # ── 형(型): 각 N형 첫 번째만 추출 ────────────────────────
    hyung_m = re.findall(r'([1-9]형\([^)]{3,40}\))', pdf_text)
    seen_h, res_h = set(), []
    for m2 in hyung_m:
        num = re.match(r'([1-9])형', m2).group(1)
        if num not in seen_h:
            seen_h.add(num); res_h.append(m2)
    result["hyung"] = " / ".join(res_h[:9]) if res_h else "해당없음"

    # ── 보통약관 + 보험기간·납입기간·연령 ──────────────────────
    # 지급사유 표에서 "보통약관" 또는 "기본계약" 행 포함 테이블 파싱
    # "보통약관(담보명)" 패턴과 보험기간/납입기간/가입나이 컬럼에서 추출
    bp_section_idx = -1
    for kw in ["1) 보통약관 및 특별약관", "1) 보통약관", "보통약관\n보 장", "나-1. 보험금"]:
        bi = pdf_text.find(kw)
        if bi >= 0:
            bp_section_idx = bi; break
    if bp_section_idx < 0:
        for kw in ["보통약관", "기본계약"]:
            bi = pdf_text.find(kw)
            if bi >= 0:
                bp_section_idx = bi; break

    if bp_section_idx >= 0:
        bp_raw = pdf_text[bp_section_idx:bp_section_idx+2000]
        # 담보명: "보통약관(담보명)" 또는 담보명 직접
        bp_name_m = re.findall(r'보통약관\(([^)]{3,50})\)', bp_raw)
        if not bp_name_m:
            # "보통약관" 다음 줄의 담보명
            for line in bp_raw.split('\n'):
                l = line.strip()
                if re.search(r'후유장해|사망|수술|입원', l) \
                   and "지급" not in l and len(l) > 4:
                    bp_name_m = [l[:50]]
                    break
        if bp_name_m:
            result["basic_policy"] = " / ".join(list(dict.fromkeys(bp_name_m))[:4])

        # 보험기간: bp_raw에서 N세만기 추출
        ip_m = list(dict.fromkeys(re.findall(r'\d+세\s*만기', bp_raw)))
        if ip_m:
            result["insurance_period"] = " / ".join(ip_m[:6])

        # 납입기간: bp_raw에서 N년납/전기납 추출
        pp_m = list(dict.fromkeys(re.findall(r'\d+년납|전기납', bp_raw)))
        if pp_m:
            result["payment_period"] = " / ".join(pp_m[:8])

        # 연령: "가입나이" 컬럼에서 최소~최대
        age_vals = re.findall(r'(\d+)세[~～](\d+)세', bp_raw)
        if age_vals:
            mins = [int(a[0]) for a in age_vals if 0 < int(a[0]) < 100]
            maxs = [int(a[1]) for a in age_vals if 0 < int(a[1]) < 100]
            if mins and maxs:
                result["age_range"] = f"만{min(mins)}세 ~ 만{max(maxs)}세"
        if result["age_range"] == "미기재":
            age_m2 = re.findall(r'(\d+)세[~～](\d+)세', bp_raw)
            if not age_m2:
                # "15세" "80세" 등 개별 값
                age_all = re.findall(r'(\d+)세', bp_raw)
                ages = [int(a) for a in age_all if 0 < int(a) < 100]
                if ages:
                    result["age_range"] = f"만{min(ages)}세 ~ 만{max(ages)}세"

    # 미추출 시 전체에서
    if result["insurance_period"] == "미기재":
        ip_all = list(dict.fromkeys(re.findall(r'\d+세\s*만기', pdf_text)))
        if ip_all: result["insurance_period"] = " / ".join(ip_all[:6])
    if result["payment_period"] == "미기재":
        pp_all = list(dict.fromkeys(re.findall(r'\d+년납|전기납', pdf_text)))
        if pp_all: result["payment_period"] = " / ".join(pp_all[:8])

    # ── 납입면제: "납입면제 사유" 섹션 원인사건 추출 ─────────────
    pw_idx = -1
    for kw in ["납입면제 사유\n", "납입면제 사유는", "납입면제에 관한 사항",
               "보험료 납입면제", "납입면제"]:
        pw_idx = pdf_text.find(kw)
        if pw_idx >= 0: break
    if pw_idx >= 0:
        pw_section = pdf_text[pw_idx:pw_idx+1500]
        pw_clean = []
        icons = "①②③④⑤⑥"
        seen_pw = set()

        def classify_pw(item):
            if "상해" in item and "80%" in item and \
               ("후유장해" in item or "장해상태" in item or "장해지급률" in item):
                return "상해80%이상후유장해", "상해80%"
            elif "질병" in item and "80%" in item and \
               ("후유장해" in item or "장해상태" in item or "장해지급률" in item):
                return "질병80%이상후유장해", "질병80%"
            elif "암" in item and "진단" in item:
                return "암진단확정", "암"
            return None, None

        # 패턴1: "- 항목" 형태 (한화 등)
        pw_raw = re.findall(r'[-]\s*([가-힣""][^\n]{3,80}(?:경우|상태|때))', pw_section)
        for item in pw_raw[:8]:
            label, key = classify_pw(item.strip())
            if label and key and key not in seen_pw:
                seen_pw.add(key)
                pw_clean.append(f"{icons[len(pw_clean)]}{label}")

        # 패턴2: 줄에서 키워드 직접 탐색 ("- 뇌졸중", "※ 뇌졸중", "「뇌졸중」" 등)
        kw_map = [
            (["뇌졸중"], "뇌졸중", "뇌졸중진단확정"),
            (["급성심근경색"], "급성심근경색증", "급성심근경색증진단확정"),
            (["특정상해성뇌출혈"], "특정상해성뇌출혈", "특정상해성뇌출혈진단확정"),
            (["상해 후유장해(80%","상해80%","상해로 장해"], "상해80%", "상해80%이상후유장해"),
            (["암(유사암 제외)", "암으로 인한", "암\" 진단", "암\"으로"], "암", "암진단확정"),
        ]
        for line in pw_section.split('\n'):
            l = line.strip()
            for kws, key, label in kw_map:
                if key not in seen_pw and any(kw in l for kw in kws) and len(pw_clean) < 6:
                    seen_pw.add(key)
                    pw_clean.append(f"{icons[len(pw_clean)]}{label}")

        if pw_clean:
            result["premium_waiver"] = " ".join(pw_clean)

    # ── 무사고전환: 요약 기재 ──────────────────────────────────
    nc_idx = -1
    for kw in ["무사고 고객 계약전환", "무사고전환", "무사고 고객"]:
        nc_idx = pdf_text.find(kw)
        if nc_idx >= 0: break
    if nc_idx >= 0:
        nc_section = pdf_text[nc_idx:nc_idx+400]
        summary = "무사고 고객 계약전환 운영"
        if "위험이 감소" in nc_section or "무사고 기간" in nc_section:
            summary += " (무사고 기간 경과 시 위험 감소에 따른 낮은 형으로 전환 가능)"
        else:
            # 첫 줄에서 핵심
            first = nc_section.split('\n')[0].strip()[:60]
            summary += f" ({first})"
        result["no_claim"] = summary
    else:
        result["no_claim"] = "해당없음"

    # ── 부가서비스: 하위 가.나.다. 제목만 ─────────────────────
    svc_idx = -1
    for kw in ["부가서비스 운영에 관한 사항", "부가서비스에 관한 사항",
               "부가서비스 운영", "부가서비스"]:
        svc_idx = pdf_text.find(kw)
        if svc_idx >= 0: break
    if svc_idx >= 0:
        svc_section = pdf_text[svc_idx:svc_idx+3000]  # 범위 확대
        svc_lines = svc_section.split('\n')
        svc_names = []
        SKIP = {"이용조건", "제공회사", "제공기간", "주요내용", "지정대리청구",
                "①", "②", "③", "서비스 내용", "운영방법", "구분"}
        for li, line in enumerate(svc_lines):
            line = line.strip()
            m_svc = re.match(r'^[가나다라마바사]\.\s*(.*)$', line)
            if m_svc:
                title = m_svc.group(1).strip()
                # 제목이 비어있으면 다음 줄과 합치기 (PDF 줄바꿈 처리)
                if not title and li+1 < len(svc_lines):
                    title = svc_lines[li+1].strip()
                if not title:
                    continue
                if not any(sk in title for sk in SKIP) and len(title) > 1:
                    # 다음 줄이 세부내용("1) 이용조건" 등)인지 확인
                    next_l = svc_lines[li+1].strip() if li+1 < len(svc_lines) else ""
                    next_l2 = svc_lines[li+2].strip() if li+2 < len(svc_lines) else ""
                    is_title = any(k in next_l for k in
                                   ["이용조건", "제공회사", "1)", "①", "주요내용"]) or \
                               any(k in next_l2 for k in
                                   ["이용조건", "제공회사", "1)", "①"])
                    if is_title or len(svc_names) == 0:
                        svc_names.append(title[:35])
                    if len(svc_names) >= 5:
                        break
        result["services"] = " / ".join(svc_names[:5]) if svc_names else "해당없음"
    else:
        result["services"] = "해당없음"

    # ── 할인: 하위 내용 요약 ────────────────────────────────────
    disc_idx = -1
    for kw in ["보험료 차등적용에 관한 사항", "보험료 차등",
               "보험료 할인에 관한 사항", "보험료할인"]:
        disc_idx = pdf_text.find(kw)
        if disc_idx >= 0: break
    if disc_idx >= 0:
        disc_section = pdf_text[disc_idx:disc_idx+500]
        if re.search(r'해당사항\s*없음|해당없음', disc_section[:120]):
            result["discounts"] = "해당없음"
        else:
            disc_items = []
            # "(N) 할인명" 패턴
            for m_d in re.finditer(
                    r'(?:\(\d+\)|[①②③])\s*([가-힣][^\n]{2,35}할인)', disc_section):
                disc_name = m_d.group(1).strip()
                # 해당 블록에서 퍼센트 추출
                block = disc_section[m_d.start():m_d.start()+250]
                rate_m = re.search(r'(\d+)%', block)
                if rate_m:
                    disc_items.append(f"{disc_name}(제1회 보험료 {rate_m.group(1)}%)")
                else:
                    disc_items.append(disc_name[:40])
            if disc_items:
                result["discounts"] = disc_items[0]  # 첫 번째만 (중복 방지)
            else:
                # 섹션 내 할인명 직접 추출
                for line in disc_section.split('\n'):
                    l = line.strip()
                    if "할인" in l and len(l) > 5 \
                       and "보험료 차등" not in l and "적용대상" not in l:
                        result["discounts"] = l[:60]
                        break
    else:
        result["discounts"] = "해당없음"

    return result

# ── A-11: 담보 추출 전체 파이프라인 ────────────────────────
def extract_coverages(file_bytes, company):
    coverages = []
    seen_keys = set()
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            in_section = False
            last_name_raw = ""
            for page in pdf.pages:
                text = page.extract_text() or ""
                if any(m in text for m in ['보험금 지급사유 및 지급금액',
                                            '보험금지급사유 및 지급금액',
                                            '보험금 지급사유 및 지급제한']):
                    in_section = True
                if not in_section: continue
                for table in page.extract_tables():
                    if not table or len(table) < 2: continue
                    header = [str(c or '').strip() for c in table[0]]
                    if company == "현대해상":
                        if not any('구 분' in h or '구분' in h for h in header): continue
                        nc, rc = 1, 2
                    elif company == "DB손해보험":
                        if '담보명' not in header: continue
                        nc = header.index('담보명')
                        rc = next((i for i,h in enumerate(header) if '지급사유' in h), nc+1)
                    else:
                        nc = next((i for i,h in enumerate(header)
                                   if re.search(r'보\s*장|담보명', h)), -1)
                        if nc < 0: continue
                        rc = next((i for i,h in enumerate(header)
                                   if re.search(r'지급\s*사유', h)), nc+1)
                    for row in table[1:]:
                        if not row or len(row) <= nc: continue
                        name_raw = str(row[nc] or '').strip()
                        reason_raw = str(row[rc] or '').strip() if len(row) > rc else ''
                        if name_raw: last_name_raw = name_raw
                        else: name_raw = last_name_raw
                        name_full = build_name(name_raw, company)
                        has_c = bool(CHEZUNG_PAT.search(name_full))
                        name_full = CHEZUNG_PAT.sub('', name_full)
                        name_full = clean_name(name_full)
                        if not is_valid(name_full): continue
                        if re.match(r'^(보 장|담보명|보장명|지급사유|구 분|No\.)', name_full): continue
                        non_v, ren_v = get_renewal_flags(name_raw)
                        reason = clean_reason(str(reason_raw).replace('\n', ' ').strip())
                        if not is_valid_reason(reason):
                            reason = guess_reason(name_full)
                        red = extract_reduction(name_full, reason_raw)
                        ex  = extract_exemption(name_full, reason_raw)
                        freq = guess_freq(name_full, reason_raw)
                        key = re.sub(r'[^\w가-힣]', '', name_full)[:25]
                        if key in seen_keys: continue
                        seen_keys.add(key)
                        coverages.append({
                            'name':       name_full,
                            'category':   categorize(name_full),
                            'has_chezung': has_c,
                            'reason':     reason[:300],
                            'freq':       freq,
                            'exemption':  ex or "없음",
                            'reduction':  red or "없음",
                            'non_renew':  non_v,
                            'renewal':    ren_v,
                        })
    except Exception as e:
        st.warning(f"{company} 담보 추출 오류: {e}")
    return coverages

# ── A-12: 검증 ──────────────────────────────────────────────
def validate_coverages(coverages):
    return {
        "①담보명 (로 시작":        sum(1 for c in coverages if re.match(r'^\(', c['name'])),
        "②담보명 숫자로 시작":     sum(1 for c in coverages if re.match(r'^\d', c['name'])),
        "③담보명 숫자로 끝남":     sum(1 for c in coverages if re.search(r'\d$', c['name'])),
        "④괄호 불일치":            sum(1 for c in coverages if c['name'].count('(')!=c['name'].count(')')),
        "⑤갱신형 텍스트 잔류":     sum(1 for c in coverages if re.search(r'\(비갱신형|\(갱신형', c['name'])),
        "⑥지급사유 감액패턴 잔류": sum(1 for c in coverages if re.search(r'\d+일\s*미만\s*\d+%|\d+년\s*미만\s*\d+%', c.get('reason',''))),
        "⑦지급사유 빈칸":          sum(1 for c in coverages if not c.get('reason')),
        "⑧진단비+수술횟수 오류":   sum(1 for c in coverages if '진단비' in c['name'] and '수술' in c.get('freq','')),
    }

# ── A-13: 매핑 ──────────────────────────────────────────────
THRESHOLD = 0.38

# 담보 카테고리별 핵심 키워드 (의미 유사도 보조)
_KEYWORD_GROUPS = [
    {"암진단", "암진단비", "암진단확정", "암보장"},
    {"유사암", "4대유사암", "기타피부암", "갑상선암", "제자리암", "경계성종양"},
    {"뇌졸중", "뇌출혈", "뇌혈관", "뇌경색"},
    {"급성심근경색", "허혈성심장", "심혈관", "심근경색"},
    {"후유장해", "80%이상후유장해", "3~100%후유장해"},
    {"수술비", "수술", "골절수술", "화상수술"},
    {"입원일당", "입원비", "입원"},
    {"중환자실", "중환자실입원"},
    {"항암방사선", "항암약물", "항암치료", "항암"},
    {"진단비", "진단"},
    {"통원", "외래", "응급실"},
    {"치료비", "통합치료비"},
    {"사망"},
    {"재활치료", "재활"},
    {"1인실", "상급병실", "1인실입원"},
    {"배상책임", "일상생활배상"},
    {"납입면제", "납입면제대상"},
]

def _keyword_sim(name_a, name_b):
    """같은 의미그룹 키워드 포함 여부로 의미 유사도 보조"""
    a_lower = name_a.replace(" ", "")
    b_lower = name_b.replace(" ", "")
    for grp in _KEYWORD_GROUPS:
        a_in = any(kw in a_lower for kw in grp)
        b_in = any(kw in b_lower for kw in grp)
        if a_in and b_in:
            return 0.2   # 같은 의미 그룹
        if a_in != b_in:
            return -0.1  # 다른 의미 그룹 (패널티)
    return 0.0

def match_score(h, c):
    """담보명 + 지급사유 텍스트 유사도 + 의미 유사도 종합"""
    h_name = re.sub(r'[\(\)（）\[\]【】]', ' ', h['name']).strip()
    c_name = re.sub(r'[\(\)（）\[\]【】]', ' ', c['name']).strip()
    # 핵심 키워드 추출 (간편/갱신형 등 제거 후)
    def core(n):
        return re.sub(r'\(간편[^)]*\)|\(갱신형\)|\[간편\]|\[갱신형\]|\(비갱신형\)', '', n).strip()
    ns = SequenceMatcher(None, core(h_name), core(c_name)).ratio()
    rs = SequenceMatcher(None, h.get('reason',''), c.get('reason','')).ratio()
    cat_bonus = 0.12 if h['category'] == c['category'] else 0.0
    kw_sim = _keyword_sim(h['name'], c['name'])
    score = ns * 0.55 + rs * 0.25 + cat_bonus + kw_sim
    return max(0.0, score)

def build_mapping(base_covs, comp_covs):
    """기준회사 담보 ↔ 비교회사 담보 1:1 매핑 (그리디)"""
    # 1단계: THRESHOLD 이상인 모든 쌍 유사도 계산
    scores = []
    for h in base_covs:
        for c in comp_covs:
            sc = match_score(h, c)
            if sc >= THRESHOLD:
                scores.append((sc, h['name'], c['name']))

    # 2단계: 유사도 내림차순 정렬 후 1:1 그리디 매핑
    scores.sort(key=lambda x: -x[0])
    used_base = set()
    used_comp = set()
    matched_pairs = {}

    comp_dict = {c['name']: c for c in comp_covs}
    for sc, h_name, c_name in scores:
        if h_name not in used_base and c_name not in used_comp:
            used_base.add(h_name)
            used_comp.add(c_name)
            matched_pairs[h_name] = comp_dict[c_name]

    # 3단계: 미매핑 기준회사 담보는 None
    return {h['name']: matched_pairs.get(h['name'], None) for h in base_covs}

# ════════════════════════════════════════════════════════════
#  B. 엑셀 빌더
# ════════════════════════════════════════════════════════════
THIN  = Side(style="thin",   color="BFBFBF")
TBOR  = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
HFILL = PatternFill("solid", fgColor="1F4E79")
H2FIL = PatternFill("solid", fgColor="2E75B6")
SFILL = PatternFill("solid", fgColor="D6E4F0")
WFILL = PatternFill("solid", fgColor="FFFFFF")
FFILL = PatternFill("solid", fgColor="FAFCFE")
HFONT = Font(name="Arial", bold=True, color="FFFFFF", size=10)
DFONT = Font(name="Arial", size=9)
BFONT = Font(name="Arial", bold=True, size=9)
CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)
LW     = Alignment(horizontal="left",   vertical="center", wrap_text=True)
CAT_C  = {"암":"FFCCCC","뇌/심":"CCF2EE","수술비":"D4E8FF",
          "입원비":"D4F0D4","치료비":"FFF0CC","통원비":"F0D4FF","기타":"EBEBEB"}
CAT_HDR= {"암":"C62828","뇌/심":"00695C","수술비":"1565C0",
          "입원비":"2E7D32","치료비":"E65100","통원비":"6A1B9A","기타":"546E7A"}
CO_C   = {"DB손해보험":"EAF4FD","KB손해보험":"EAF8EA","메리츠화재":"FFFAE0",
          "삼성화재":"FDEAEE","한화손해보험":"F5EAF5","현대해상":"E0F5F2",
          "흥국화재":"FFF9E6","롯데손해보험":"F0F0FF"}

def mkf(hex_color):
    return PatternFill("solid", fgColor=hex_color)

def Hc(cell, val, fill=None, font=None, align=CENTER):
    cell.value = str(val) if val else ""
    cell.fill = fill or HFILL; cell.font = font or HFONT
    cell.alignment = align; cell.border = TBOR

def Dc(cell, val, fill=None, font=None, align=LW):
    cell.value = str(val) if val else ""
    cell.fill = fill or FFILL; cell.font = font or DFONT
    cell.alignment = align; cell.border = TBOR

def Tc(ws, r, c1, c2, val, fill=None, font=None):
    ws.merge_cells(f"{get_column_letter(c1)}{r}:{get_column_letter(c2)}{r}")
    c = ws.cell(r, c1); c.value = str(val)
    c.fill = fill or HFILL; c.font = font or HFONT
    c.alignment = CENTER; c.border = TBOR

def cs(s):
    if not s: return ""
    s = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', str(s))
    return re.sub(r'\s+', ' ', s).strip()

# ── 시트: 상품구조 ───────────────────────────────────────────
def build_sheet_structure(wb, structures):
    ws = wb.create_sheet("상품구조")
    ws.sheet_view.showGridLines = False
    cos = list(structures.keys())
    total_cols = 1 + len(cos)
    Tc(ws, 1, 1, total_cols, "3N5 손해보험 간편건강보험 상품구조 비교표",
       fill=mkf("EBF3FB"), font=Font(name="Arial",bold=True,size=14,color="1F4E79"))
    ws.row_dimensions[1].height = 32
    # 헤더
    ws.row_dimensions[2].height = 26
    Hc(ws.cell(2,1), "항목")
    for i, co in enumerate(cos, 2):
        Hc(ws.cell(2, i), co, fill=mkf(CO_C.get(co,"EEEEEE")),
           font=Font(name="Arial",bold=True,size=10,color="333333"))
    ITEMS = [
        ("product_name",     "상품명"),
        ("jong",             "종(種)"),
        ("hyung",            "형(型)"),
        ("basic_policy",     "보통약관"),
        ("insurance_period", "보험기간"),
        ("payment_period",   "납입기간"),
        ("age_range",        "연령"),
        ("premium_waiver",   "납입면제"),
        ("no_claim",         "무사고전환"),
        ("services",         "부가서비스"),
        ("discounts",        "할인"),
    ]
    for r, (key, label) in enumerate(ITEMS, 3):
        ws.row_dimensions[r].height = 50
        fill = SFILL if r%2==0 else WFILL
        Dc(ws.cell(r,1), label, fill=mkf("2E75B6"),
           font=Font(name="Arial",bold=True,color="FFFFFF",size=9), align=CENTER)
        ws.cell(r,1).border = TBOR
        for i, co in enumerate(cos, 2):
            val = structures[co].get(key, "미기재")
            Dc(ws.cell(r,i), cs(val), fill=fill)
    ws.column_dimensions["A"].width = 14
    for i in range(2, total_cols+1):
        ws.column_dimensions[get_column_letter(i)].width = 40
    ws.freeze_panes = "B3"
    return ws

# ── 시트: 회사별 담보 ────────────────────────────────────────
def build_sheet_company(wb, company, coverages):
    ws = wb.create_sheet(company)
    ws.sheet_view.showGridLines = False
    Tc(ws,1,1,9, f"{company} 담보목록 ({len(coverages)}개)",
       fill=mkf("EBF3FB"), font=Font(name="Arial",bold=True,size=12,color="1F4E79"))
    ws.row_dimensions[1].height = 26
    COLS = ["No.","담보명","분류","지급사유","지급횟수","면책기간","감액내용","비갱신형","갱신형"]
    ws.row_dimensions[2].height = 22
    for i,h in enumerate(COLS,1): Hc(ws.cell(2,i), h)
    co_fill = mkf(CO_C.get(company,"FAFCFE"))
    for r, cov in enumerate(coverages, 3):
        cat = cov.get("category","")
        cat_f = mkf(CAT_C.get(cat,"EBEBEB"))
        ws.row_dimensions[r].height = 38
        vals = [r-2, cov['name'], cat, cov.get('reason',''),
                cov.get('freq',''), cov.get('exemption','없음'),
                cov.get('reduction','없음'),
                cov.get('non_renew',''), cov.get('renewal','')]
        for i,v in enumerate(vals,1):
            c = ws.cell(r,i)
            if i==1:   Dc(c,v,fill=co_fill,align=CENTER)
            elif i==3: Dc(c,v,fill=cat_f,font=BFONT,align=CENTER)
            elif i in(8,9): Dc(c,v,align=CENTER)
            else: Dc(c,v)
    for i,w in enumerate([6,44,10,70,16,10,22,8,8],1):
        ws.column_dimensions[get_column_letter(i)].width=w
    ws.freeze_panes="A3"
    return ws

# ── 시트: 담보목록_전체 ─────────────────────────────────────
def build_sheet_all(wb, all_covs_dict):
    ws = wb.create_sheet("담보목록_전체")
    ws.sheet_view.showGridLines = False
    total = sum(len(v) for v in all_covs_dict.values())
    Tc(ws,1,1,9,f"3N5 손해보험 담보목록 전체 (총 {total}개)",
       fill=mkf("EBF3FB"),font=Font(name="Arial",bold=True,size=13,color="1F4E79"))
    ws.row_dimensions[1].height=28
    COLS=["회사명","담보명","분류","지급사유","지급횟수","면책기간","감액내용","비갱신형","갱신형"]
    ws.row_dimensions[2].height=24
    for i,h in enumerate(COLS,1): Hc(ws.cell(2,i),h)
    rn=3
    for co, covs in all_covs_dict.items():
        Tc(ws,rn,1,9,f"[{co}]  {len(covs)}개 담보",fill=H2FIL)
        ws.row_dimensions[rn].height=20; rn+=1
        co_f=mkf(CO_C.get(co,"FAFCFE"))
        for cov in covs:
            cat=cov.get("category","")
            cat_f=mkf(CAT_C.get(cat,"EBEBEB"))
            ws.row_dimensions[rn].height=38
            for i,v in enumerate([co,cov['name'],cat,cov.get('reason',''),
                cov.get('freq',''),cov.get('exemption','없음'),
                cov.get('reduction','없음'),cov.get('non_renew',''),cov.get('renewal','')],1):
                c=ws.cell(rn,i)
                if i==1: Dc(c,v,fill=co_f,font=BFONT,align=CENTER)
                elif i==3: Dc(c,v,fill=cat_f,font=BFONT,align=CENTER)
                elif i in(8,9): Dc(c,v,align=CENTER)
                else: Dc(c,v)
            rn+=1
    for i,w in enumerate([13,44,10,70,16,10,22,8,8],1):
        ws.column_dimensions[get_column_letter(i)].width=w
    ws.freeze_panes="A3"
    return ws

# ── 시트: 회사별_담보수_요약 ────────────────────────────────
def build_sheet_summary(wb, all_covs_dict):
    ws = wb.create_sheet("회사별_담보수_요약")
    ws.sheet_view.showGridLines=False
    CATS=["암","뇌/심","수술비","입원비","치료비","통원비","기타"]
    Tc(ws,1,1,9,"회사별 담보 분류 요약",
       fill=mkf("EBF3FB"),font=Font(name="Arial",bold=True,size=13,color="1F4E79"))
    ws.row_dimensions[1].height=28
    COLS=["회사명","총담보수"]+CATS
    ws.row_dimensions[2].height=24
    for i,h in enumerate(COLS,1):
        c=ws.cell(2,i)
        if i<=2: Hc(c,h)
        else: Hc(c,h,fill=mkf(CAT_C.get(CATS[i-3],"EBEBEB")),
                 font=Font(name="Arial",bold=True,size=9,color="333333"))
    cos = list(all_covs_dict.keys())
    for r,co in enumerate(cos,3):
        covs=all_covs_dict[co]
        cnt=Counter(c["category"] for c in covs)
        vals=[co,len(covs)]+[cnt.get(cat,0) for cat in CATS]
        fill=SFILL if r%2==0 else WFILL
        ws.row_dimensions[r].height=22
        for i,v in enumerate(vals,1):
            Dc(ws.cell(r,i),str(v),fill=fill)
            ws.cell(r,i).alignment=CENTER
            if i==1: ws.cell(r,i).font=BFONT
    total_r=len(cos)+3
    all_cnt=Counter()
    for covs in all_covs_dict.values():
        for cov in covs: all_cnt[cov["category"]]+=1
    sv=["합계",str(sum(len(v) for v in all_covs_dict.values()))]+[str(all_cnt.get(cat,0)) for cat in CATS]
    ws.row_dimensions[total_r].height=24
    for i,v in enumerate(sv,1):
        c=ws.cell(total_r,i,v); c.fill=H2FIL
        c.font=Font(name="Arial",bold=True,size=10,color="FFFFFF")
        c.alignment=CENTER; c.border=TBOR
    for i,w in enumerate([14,10]+[12]*7,1):
        ws.column_dimensions[get_column_letter(i)].width=w
    ws.freeze_panes="A3"
    return ws

# ── 시트: 담보비교표 ────────────────────────────────────────
def build_sheet_comparison(wb, base_co, base_covs, compare_dict, mappings):
    """
    base_covs: 기준회사 담보 리스트 (순서 유지)
    compare_dict: {co: [cov_list]}
    mappings: {co: {h_name: matched_cov or None}}
    """
    ws = wb.create_sheet("담보비교표")
    ws.sheet_view.showGridLines = False
    COMPARE = list(compare_dict.keys())
    ALL_COS = [base_co] + COMPARE
    CO_START = {}
    col = 2  # A=분류
    for co in ALL_COS:
        CO_START[co] = col
        col += 5
    TOTAL_COLS = col - 1

    # 타이틀
    Tc(ws,1,1,TOTAL_COLS,f"3N5 손해보험 담보 비교표  (기준: {base_co})",
       fill=mkf("1F4E79"),font=Font(name="Arial",bold=True,size=14,color="FFFFFF"))
    ws.row_dimensions[1].height=34
    # 분류 헤더
    ws.row_dimensions[2].height=26
    Hc(ws.cell(2,1),"분류")
    for co in ALL_COS:
        st2=CO_START[co]
        ws.merge_cells(f"{get_column_letter(st2)}2:{get_column_letter(st2+4)}2")
        c=ws.cell(2,st2); c.value=co
        c.fill=mkf(CO_C.get(co,"EEEEEE"))
        c.font=Font(name="Arial",bold=True,size=11,color="333333")
        c.alignment=CENTER; c.border=TBOR
    # 열 헤더
    ws.row_dimensions[3].height=22
    Hc(ws.cell(3,1),"분류")
    for co in ALL_COS:
        st2=CO_START[co]
        for j,sh in enumerate(["담보명","지급사유","지급횟수","면책기간","감액내용"]):
            c=ws.cell(3,st2+j)
            c.value=sh; c.fill=mkf("3D85C8")
            c.font=Font(name="Arial",bold=True,size=9,color="FFFFFF")
            c.alignment=CENTER; c.border=TBOR
    # 열 너비
    ws.column_dimensions["A"].width=8
    for co in ALL_COS:
        st2=CO_START[co]
        for j,w in enumerate([32,50,16,10,18]):
            ws.column_dimensions[get_column_letter(st2+j)].width=w

    def write_row(row, cat, base_cov, comp_results):
        ws.row_dimensions[row].height=45
        # 분류
        Dc(ws.cell(row,1), cat, fill=mkf(CAT_C.get(cat,"EBEBEB")),
           font=Font(name="Arial",bold=True,size=9), align=CENTER)
        # 기준회사
        st2=CO_START[base_co]
        if base_cov:
            bg=CO_C.get(base_co,"FAFCFE")
            for j,(key,fnt) in enumerate(zip(
                ['name','reason','freq','exemption','reduction'],
                [BFONT,DFONT,DFONT,DFONT,DFONT])):
                Dc(ws.cell(row,st2+j), cs(base_cov.get(key,'')),
                   fill=mkf(bg), font=fnt)
        else:
            for j in range(5): Dc(ws.cell(row,st2+j),"",fill=mkf("F8F8F8"))
        # 비교회사
        for co in COMPARE:
            st2=CO_START[co]
            m=comp_results.get(co)
            if m:
                bg=CO_C.get(co,"FAFCFE")
                for j,(key,fnt) in enumerate(zip(
                    ['name','reason','freq','exemption','reduction'],
                    [BFONT,DFONT,DFONT,DFONT,DFONT])):
                    Dc(ws.cell(row,st2+j), cs(m.get(key,'')), fill=mkf(bg), font=fnt)
            else:
                c=ws.cell(row,st2)
                c.value="해당없음"; c.fill=mkf("F0F0F0")
                c.font=Font(name="Arial",size=9,color="999999",italic=True)
                c.alignment=CENTER; c.border=TBOR
                for j in range(1,5): Dc(ws.cell(row,st2+j),"",fill=mkf("F0F0F0"))

    cur_row=4; cur_cat=None
    for h in base_covs:
        cat=h.get("category","기타")
        if cat!=cur_cat:
            ws.merge_cells(f"A{cur_row}:{get_column_letter(TOTAL_COLS)}{cur_row}")
            c=ws.cell(cur_row,1)
            c.value=f"▶ {cat}"; c.fill=mkf(CAT_HDR.get(cat,"546E7A"))
            c.font=Font(name="Arial",bold=True,size=10,color="FFFFFF")
            c.alignment=LW; c.border=TBOR
            ws.row_dimensions[cur_row].height=20
            cur_row+=1; cur_cat=cat
        comp_results={co: mappings.get(co,{}).get(h['name']) for co in COMPARE}
        write_row(cur_row, cat, h, comp_results)
        cur_row+=1

    # 비교사 단독 담보
    used={co:set() for co in COMPARE}
    for co in COMPARE:
        for m in mappings.get(co,{}).values():
            if m: used[co].add(m['name'])
    extras={co:[c for c in compare_dict[co] if c['name'] not in used[co]] for co in COMPARE}
    if any(extras.values()):
        ws.merge_cells(f"A{cur_row}:{get_column_letter(TOTAL_COLS)}{cur_row}")
        c=ws.cell(cur_row,1); c.value="▶ 비교사 단독 담보 (기준회사 미보유)"
        c.fill=mkf("37474F"); c.font=Font(name="Arial",bold=True,size=11,color="FFFFFF")
        c.alignment=LW; c.border=TBOR; ws.row_dimensions[cur_row].height=22; cur_row+=1
        for co in COMPARE:
            if not extras[co]: continue
            ws.merge_cells(f"A{cur_row}:{get_column_letter(TOTAL_COLS)}{cur_row}")
            c=ws.cell(cur_row,1); c.value=f"  [{co}] 단독 담보 ({len(extras[co])}개)"
            c.fill=mkf(CO_C.get(co,"EEEEEE")); c.font=Font(name="Arial",bold=True,size=10,color="333333")
            c.alignment=LW; c.border=TBOR; ws.row_dimensions[cur_row].height=20; cur_row+=1
            ec=None
            for e in sorted(extras[co],key=lambda x:(["암","뇌/심","수술비","입원비","치료비","통원비","기타"].index(x.get("category","기타")) if x.get("category","기타") in ["암","뇌/심","수술비","입원비","치료비","통원비","기타"] else 99, x['name'])):
                if e.get("category","")!=ec:
                    ec=e.get("category","")
                    ws.merge_cells(f"A{cur_row}:{get_column_letter(TOTAL_COLS)}{cur_row}")
                    ct=ws.cell(cur_row,1); ct.value=f"▶ {ec}"; ct.fill=mkf(CAT_HDR.get(ec,"546E7A"))
                    ct.font=Font(name="Arial",bold=True,size=10,color="FFFFFF")
                    ct.alignment=LW; ct.border=TBOR; ws.row_dimensions[cur_row].height=20; cur_row+=1
                comp_r={c2:(e if c2==co else None) for c2 in COMPARE}
                write_row(cur_row, e.get("category","기타"), None, comp_r)
                cur_row+=1
    ws.freeze_panes="B4"
    return ws

# ── 전체 엑셀 빌드 ──────────────────────────────────────────
def build_full_excel(structures, all_covs_dict, base_co=None):
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    # 시트 순서: 상품구조 → 담보목록_전체 → 요약 → 회사별 → 담보비교표
    build_sheet_structure(wb, structures)
    build_sheet_all(wb, all_covs_dict)
    build_sheet_summary(wb, all_covs_dict)
    cos = list(all_covs_dict.keys())
    for co in cos:
        build_sheet_company(wb, co, all_covs_dict[co])
    # 담보비교표: 한화 기준 (없으면 첫 번째 회사)
    if base_co and base_co in all_covs_dict and len(cos) > 1:
        compare_cos = [c for c in cos if c != base_co]
        compare_dict = {c: all_covs_dict[c] for c in compare_cos}
        mappings = {}
        for co in compare_cos:
            mappings[co] = build_mapping(all_covs_dict[base_co], all_covs_dict[co])
        build_sheet_comparison(wb, base_co, all_covs_dict[base_co], compare_dict, mappings)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()

# ── Word 보고서 저장 ─────────────────────────────────────────
def save_word_report(report_text, base_co, comp_cos):
    try:
        from docx import Document
        from docx.shared import RGBColor, Pt
        doc = Document()
        title = doc.add_heading("한화손해보험 상품 경쟁력 강화 방안 보고서", 0)
        if title.runs: title.runs[0].font.color.rgb = RGBColor(0xFF,0x6B,0x00)
        doc.add_paragraph(f"기준: {base_co} vs {', '.join(comp_cos)}")
        for line in report_text.split('\n'):
            if line.startswith('### '): doc.add_heading(line[4:], 2)
            elif line.startswith('## '): doc.add_heading(line[3:], 1)
            elif line.startswith('- '): doc.add_paragraph(line[2:], style='List Bullet')
            elif line.strip(): doc.add_paragraph(line)
        buf = io.BytesIO(); doc.save(buf)
        return buf.getvalue()
    except ImportError:
        return None

# ════════════════════════════════════════════════════════════
#  C. session_state 초기화
# ════════════════════════════════════════════════════════════
# dict 타입 키: 없거나 잘못된 타입이면 초기화
for key in ['pdf_bytes','pdf_texts','pdf_structures','pdf_coverages',
            'tab1_results','tab2_results','pdf_file_ids']:
    if key not in st.session_state or not isinstance(st.session_state[key], dict):
        st.session_state[key] = {}

# tab3_report는 None 또는 dict
if 'tab3_report' not in st.session_state:
    st.session_state['tab3_report'] = None

if 'current_menu' not in st.session_state:
    st.session_state.current_menu = "파일 업로드"

# ════════════════════════════════════════════════════════════
#  D. 사이드바: 로고 + 회사 목록 + 메뉴
# ════════════════════════════════════════════════════════════
companies_loaded = list(st.session_state.pdf_coverages.keys())

with st.sidebar:
    # 로고
    st.markdown("""
    <div style="background:#FF6B00;padding:14px 16px;border-radius:8px;margin-bottom:16px">
      <div style="font-size:14px;font-weight:700;color:#fff">H &nbsp; InsureCompare</div>
      <div style="font-size:11px;color:rgba(255,255,255,.8)">한화손보 상품분석시스템</div>
    </div>
    """, unsafe_allow_html=True)

    # 업로드된 회사 목록
    if companies_loaded:
        st.markdown('<div style="font-size:11px;color:#999;font-weight:600;margin-bottom:6px">회사</div>',
                    unsafe_allow_html=True)
        base_co_sidebar = companies_loaded[0]
        for co in companies_loaded:
            is_base = (co == base_co_sidebar)
            dot = '<span style="color:#FF6B00">●</span>' if is_base else '<span style="color:#ccc">●</span>'
            badge = ' <span style="background:#FF6B00;color:#fff;font-size:9px;padding:1px 6px;border-radius:8px">기준</span>' if is_base else ''
            st.markdown(
                f'<div style="padding:3px 4px;font-size:13px">{dot} {co}{badge}</div>',
                unsafe_allow_html=True)
        st.markdown('<hr style="border:none;border-top:1px solid #eee;margin:10px 0">',
                    unsafe_allow_html=True)

    # 메뉴 버튼
    st.markdown('<div style="font-size:11px;color:#999;font-weight:600;margin-bottom:6px">메뉴</div>',
                unsafe_allow_html=True)
    MENUS = [
        ("📁 파일 업로드",   "파일 업로드"),
        ("🗂 상품 구조 비교", "상품 구조 비교"),
        ("🔍 담보 내용 비교", "담보 내용 비교"),
        ("📋 경쟁력 강화방안",  "경쟁력 강화방안"),
    ]
    for label, key in MENUS:
        is_active = st.session_state.current_menu == key
        if st.button(label, key=f"menu_{key}", use_container_width=True,
                     type="primary" if is_active else "secondary"):
            st.session_state.current_menu = key

    # 전체 엑셀 생성 (하단)
    st.markdown('<hr style="border:none;border-top:1px solid #eee;margin:14px 0">',
                unsafe_allow_html=True)
    st.markdown('<div style="font-size:11px;color:#999;font-weight:600;margin-bottom:6px">내보내기</div>',
                unsafe_allow_html=True)
    if st.button("🗂️ 전체 엑셀 생성", key="full_excel", use_container_width=True):
        if not companies_loaded:
            st.error("PDF를 먼저 업로드하세요.")
        else:
            with st.spinner("⏳ 전체 엑셀 생성 중..."):
                all_structs = {co: st.session_state.pdf_structures.get(co,{}) for co in companies_loaded}
                all_covs    = {co: st.session_state.pdf_coverages.get(co,[])  for co in companies_loaded}
                base = "한화손해보험" if "한화손해보험" in companies_loaded else companies_loaded[0]
                excel_bytes = build_full_excel(all_structs, all_covs, base_co=base)
            st.download_button(
                "⬇️ 다운로드",
                data=excel_bytes,
                file_name="insurance_3N5_comparison_v7.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="dl_full_excel"
            )

# ════════════════════════════════════════════════════════════
#  E. 파일 업로드 페이지
# ════════════════════════════════════════════════════════════
UPLOAD_COMPANIES = ["한화손해보험","삼성화재","현대해상","메리츠화재","KB손해보험","DB손해보험"]

def _parse_uploaded_file(f, co):
    """업로드된 파일 파싱 후 session_state 저장
    - raw bytes는 저장하지 않음 (메모리 절약)
    - 파싱을 최대한 빠르게 처리 (Render 30초 타임아웃 대응)
    """
    # 이미 파싱된 경우 건너뜀 (on_change 콜백으로만 호출되므로 단순 체크로 충분)
    if co in st.session_state.pdf_coverages and st.session_state.pdf_coverages[co]:
        return
    try:
        f.seek(0)
        raw = f.read()
        if not raw:
            st.error(f"{co}: 파일이 비어있습니다.")
            return
        file_size_mb = len(raw) / 1024 / 1024
        if file_size_mb > 30:
            st.error(f"{co}: 파일이 너무 큽니다 ({file_size_mb:.1f}MB). 30MB 이하 파일만 가능합니다.")
            return

        progress = st.progress(0, text=f"⏳ {co} 파싱 중...")

        # ── 1단계: 텍스트 추출 (빠른 모드) ──────────────────────
        progress.progress(10, text=f"📄 {co} 텍스트 추출 중...")
        txt_pages = []
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            total_pages = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                try:
                    # 텍스트 추출 전략: bbox 최적화로 속도 향상
                    txt_pages.append(page.extract_text(
                        x_tolerance=3, y_tolerance=3
                    ) or "")
                except Exception:
                    txt_pages.append("")
                if i % 5 == 0:
                    pct = 10 + int((i / max(total_pages, 1)) * 30)
                    progress.progress(pct, text=f"📄 {co} 텍스트 추출 중... ({i+1}/{total_pages}페이지)")
        txt = "\n".join(txt_pages)

        # ── 2단계: 상품구조 추출 ─────────────────────────────────
        progress.progress(45, text=f"🗂 {co} 상품구조 분석 중...")
        structure = extract_product_structure(txt, co)

        # ── 3단계: 담보 추출 (핵심 최적화: 지급사유 섹션 이후만) ──
        progress.progress(60, text=f"🔍 {co} 담보 추출 중...")
        covs = extract_coverages(raw, co)

        # ── 4단계: session_state 저장 ────────────────────────────
        progress.progress(95, text=f"💾 {co} 저장 중...")
        st.session_state.pdf_texts[co]      = txt
        st.session_state.pdf_structures[co] = structure
        st.session_state.pdf_coverages[co]  = covs
        st.session_state.pdf_bytes[co]      = True  # 완료 플래그

        progress.progress(100, text=f"✅ {co} 완료!")
        progress.empty()
        st.success(f"✅ {co}: {len(covs)}개 담보 추출 완료 ({file_size_mb:.1f}MB, {total_pages}페이지)")

    except Exception as e:
        st.session_state.pdf_bytes.pop(co, None)
        st.session_state.pdf_texts.pop(co, None)
        st.session_state.pdf_structures.pop(co, None)
        st.session_state.pdf_coverages.pop(co, None)
        st.error(f"❌ {co} 파싱 오류: {e}")
        import traceback
        st.text(traceback.format_exc()[:500])

def show_upload_page():
    st.markdown("## 📁 파일 업로드")

    # 업로드된 회사 수 표시
    uploaded_count = len([co for co in UPLOAD_COMPANIES if co in st.session_state.pdf_bytes])
    if uploaded_count > 0:
        st.info(f"✅ 현재 **{uploaded_count}개** 회사 요약서가 업로드되어 있습니다. 좌측 메뉴에서 분석을 실행하세요.")

    # ── 콜백 헬퍼: 특정 회사의 PDF 파싱 트리거 ──────────────────
    def _make_upload_callback(company):
        """file_uploader on_change 콜백 팩토리.
        Streamlit은 어떤 위젯이 바뀌어도 전체를 리렌더링하므로,
        on_change 콜백 안에서만 파싱을 실행해 다른 회사 재파싱을 방지한다."""
        def _cb():
            f = st.session_state.get(f"up_{company}")
            if f is None:
                return
            # 기존 데이터 초기화 (파일 교체 대응)
            for k in ['pdf_bytes', 'pdf_texts', 'pdf_structures',
                      'pdf_coverages', 'pdf_file_ids']:
                st.session_state[k].pop(company, None)
            _parse_uploaded_file(f, company)
        return _cb

    # 기준사 카드 (전폭)
    base_co_up = UPLOAD_COMPANIES[0]
    is_uploaded_base = base_co_up in st.session_state.pdf_bytes
    dot_base = "🟠" if is_uploaded_base else "⚪"
    badge_html = '<span style="background:#FF6B00;color:#fff;font-size:11px;padding:2px 8px;border-radius:8px;margin-left:6px">기준사</span>'
    card_style = "border:1.5px solid #FF6B00;border-radius:8px;padding:16px;background:#FFF8F3;margin-bottom:16px"
    with st.container():
        st.markdown(f'<div style="{card_style}"><span style="font-size:15px;font-weight:600">{dot_base} {base_co_up}</span>{badge_html}</div>',
                    unsafe_allow_html=True)
        if is_uploaded_base:
            st.markdown('<div style="color:#2E7D32;font-size:12px;margin-bottom:8px">✅ 업로드 완료 (다른 파일로 교체하려면 아래에 업로드)</div>',
                        unsafe_allow_html=True)
        else:
            st.markdown('<div style="color:#999;font-size:12px;margin-bottom:8px">파일을 추가하세요</div>',
                        unsafe_allow_html=True)
        st.file_uploader("", type="pdf", key=f"up_{base_co_up}",
                         label_visibility="collapsed",
                         on_change=_make_upload_callback(base_co_up))

    # 나머지 5개 회사 3열 그리드
    others = UPLOAD_COMPANIES[1:]
    for row_start in range(0, len(others), 3):
        cols = st.columns(3)
        for ci, co in enumerate(others[row_start:row_start+3]):
            with cols[ci]:
                is_up = co in st.session_state.pdf_bytes
                dot = "🟠" if is_up else "⚪"
                c_style = f"border:1px solid {'#FF6B00' if is_up else '#eee'};border-radius:8px;padding:14px;background:{'#FFF8F3' if is_up else '#fff'}"
                st.markdown(f'<div style="{c_style}"><span style="font-size:13px;font-weight:600">{dot} {co}</span></div>',
                            unsafe_allow_html=True)
                if is_up:
                    st.markdown('<div style="color:#2E7D32;font-size:11px;margin:4px 0 6px 2px">✅ 업로드 완료</div>',
                                unsafe_allow_html=True)
                else:
                    st.markdown('<div style="color:#999;font-size:11px;margin:4px 0 6px 2px">파일을 추가하세요</div>',
                                unsafe_allow_html=True)
                st.file_uploader("", type="pdf", key=f"up_{co}",
                                 label_visibility="collapsed",
                                 on_change=_make_upload_callback(co))

    # 기능 안내
    info_cols = st.columns(3)
    with info_cols[0]:
        st.markdown("""**🗂 상품구조**
운영형태·가입나이·보험기간·납입기간·납입면제·무사고전환·할인·부가서비스 (2회 검증)""")
    with info_cols[1]:
        st.markdown("""**🔍 담보비교**
담보명 유사도 → 지급사유 최종매칭 · 기호·공백 제외""")
    with info_cols[2]:
        st.markdown("""**📋 경쟁력 강화방안**
구분·기준·비교·시사점 4컬럼 표·구조·담보 비교 실행 필요""")
    st.info("ⓘ 업로드 시 텍스트만 추출합니다. 각 분석 메뉴에서 요약서를 선택하고 실행 버튼을 누르세요.")

# ════════════════════════════════════════════════════════════
#  F. 메뉴별 콘텐츠 함수
# ════════════════════════════════════════════════════════════
def show_tab1_content():
    companies_loaded_fn = list(st.session_state.pdf_coverages.keys())
    st.markdown("### 🗂 상품 구조 비교")
    if not companies_loaded_fn:
        st.warning("📂 업로드된 PDF가 없습니다. 파일 업로드 메뉴에서 요약서를 먼저 업로드해주세요.")
        if st.button("📁 파일 업로드 메뉴로 이동", type="primary"):
            st.session_state.current_menu = "파일 업로드"
        return
    col1, col2 = st.columns(2)
    with col1:
        base_co1 = st.selectbox("기준회사", companies_loaded_fn, key="t1_base")
    with col2:
        comp_cos1 = st.multiselect(
            "비교회사 (선택 안 하면 기준회사만)",
            [c for c in companies_loaded_fn if c != base_co1], key="t1_compare"
        )
    target_cos1 = [base_co1] + comp_cos1

    if st.button("📋 상품구조 분석 실행", key="t1_run", type="primary"):
        results1 = {co: st.session_state.pdf_structures.get(co, {}) for co in target_cos1}
        st.session_state.tab1_results = results1

    if st.session_state.tab1_results:
        results1 = st.session_state.tab1_results
        cos1 = list(results1.keys())
        ITEMS1 = [
            ("product_name",     "상품명"),
            ("jong",             "종(種)"),
            ("hyung",            "형(型)"),
            ("basic_policy",     "보통약관"),
            ("insurance_period", "보험기간"),
            ("payment_period",   "납입기간"),
            ("age_range",        "연령"),
            ("premium_waiver",   "납입면제"),
            ("no_claim",         "무사고전환"),
            ("services",         "부가서비스"),
            ("discounts",        "할인"),
        ]
        header_html = "".join(
            f"<th style='background:#FF6B00;color:#fff;padding:8px 12px;text-align:left;font-size:13px'>{co}</th>"
            for co in cos1)
        rows_html = ""
        for key, label in ITEMS1:
            vals = [results1[co].get(key, "미기재") for co in cos1]
            is_diff = len(set(str(v) for v in vals)) > 1 and len(cos1) > 1
            row_bg = "background:rgba(255,107,0,0.06);" if is_diff else ""
            diff_mark = " <span style='color:#FF6B00;font-size:10px;font-weight:600'>◀ 차이</span>" if is_diff else ""
            cells = "".join(
                f"<td style='padding:7px 12px;font-size:12px;vertical-align:top;border-bottom:0.5px solid #eee'>{str(v)[:200]}{diff_mark if i==0 and is_diff else ''}</td>"
                for i, v in enumerate(vals)
            )
            rows_html += f"""<tr style='{row_bg}'>
              <td style='padding:7px 12px;font-size:12px;font-weight:600;color:#FF6B00;
                  background:#FFF8F3;border-bottom:0.5px solid #eee;border-right:1px solid #FFD4A8;
                  white-space:nowrap;min-width:90px'>{label}</td>{cells}</tr>"""
        table_html = f"""
        <div style='border:1px solid #eee;border-radius:8px;overflow:hidden;font-family:Arial,sans-serif'>
          <table style='width:100%;border-collapse:collapse'>
            <thead><tr>
              <th style='background:#333;color:#fff;padding:8px 12px;text-align:left;font-size:13px;min-width:90px'>항목</th>
              {header_html}
            </tr></thead>
            <tbody>{rows_html}</tbody>
          </table>
        </div>
        <div style='font-size:11px;color:#888;margin-top:4px'>※ 주황 음영 행: 회사 간 차이가 있는 항목</div>
        """
        st.components.v1.html(table_html, height=min(80 + len(ITEMS1)*55, 700), scrolling=False)
        wb_t1 = openpyxl.Workbook(); wb_t1.remove(wb_t1.active)
        build_sheet_structure(wb_t1, {co: results1[co] for co in cos1})
        buf1 = io.BytesIO(); wb_t1.save(buf1)
        st.download_button(
            "⬇️ 엑셀 다운로드 (상품구조)",
            data=buf1.getvalue(),
            file_name="상품구조_비교.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )


def show_tab2_content():
    companies_loaded_fn = list(st.session_state.pdf_coverages.keys())
    st.markdown("### 🔍 담보 내용 비교")
    if not companies_loaded_fn:
        st.warning("📂 업로드된 PDF가 없습니다. 파일 업로드 메뉴에서 요약서를 먼저 업로드해주세요.")
        if st.button("📁 파일 업로드 메뉴로 이동", type="primary"):
            st.session_state.current_menu = "파일 업로드"
        return
    col1, col2 = st.columns(2)
    with col1:
        base_co2 = st.selectbox("기준회사", companies_loaded_fn, key="t2_base")
    with col2:
        comp_cos2 = st.multiselect(
            "비교회사 (선택 안 하면 기준회사만 분석)",
            [c for c in companies_loaded_fn if c != base_co2], key="t2_compare"
        )

    if st.button("📊 담보 분석 실행", key="t2_run", type="primary"):
        st.session_state.tab2_results = {}
        base_covs2 = st.session_state.pdf_coverages.get(base_co2, [])
        if not comp_cos2:
            st.session_state.tab2_results = {
                "mode": "single", "base": base_co2,
                "base_covs": base_covs2, "compare": {}
            }
        else:
            mappings2 = {}
            compare_covs2 = {}
            for co in comp_cos2:
                c_covs = st.session_state.pdf_coverages.get(co, [])
                compare_covs2[co] = c_covs
                with st.spinner(f"🔄 {co} 매핑 중..."):
                    mappings2[co] = build_mapping(base_covs2, c_covs)
            st.session_state.tab2_results = {
                "mode": "compare", "base": base_co2,
                "base_covs": base_covs2,
                "compare": compare_covs2,
                "mappings": mappings2
            }

    if st.session_state.tab2_results:
        r2 = st.session_state.tab2_results
        mode2 = r2.get("mode","single")
        b2    = r2.get("base","")
        bcovs2= r2.get("base_covs",[])

        if mode2 == "single":
            st.markdown(f"**{b2}** 담보 목록 ({len(bcovs2)}개)")
            checks2 = validate_coverages(bcovs2)
            with st.expander("🔍 검증 결과"):
                for k,v in checks2.items():
                    st.markdown(f"{'✅' if v==0 else '❌'} {k}: {v}개")
            cats_all = sorted(set(c.get("category","기타") for c in bcovs2))
            sel_cats = st.multiselect("분류 필터", cats_all, default=cats_all, key="t2_cat_filter")
            filtered = [c for c in bcovs2 if c.get("category","기타") in sel_cats]
            import pandas as pd
            df2 = pd.DataFrame([{
                "No.": i+1, "담보명": c['name'], "분류": c.get("category",""),
                "지급사유": c.get("reason",""), "지급횟수": c.get("freq",""),
                "면책기간": c.get("exemption","없음"), "감액내용": c.get("reduction","없음"),
                "비갱신형": c.get("non_renew",""), "갱신형": c.get("renewal",""),
            } for i,c in enumerate(filtered)])
            st.dataframe(df2, use_container_width=True, height=500)
            wb_t2 = openpyxl.Workbook(); wb_t2.remove(wb_t2.active)
            build_sheet_company(wb_t2, b2, bcovs2)
            buf2 = io.BytesIO(); wb_t2.save(buf2)
            st.download_button(
                f"⬇️ 엑셀 다운로드 ({b2} 담보목록)",
                data=buf2.getvalue(), file_name=f"{b2}_담보목록.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        else:
            comp2     = r2.get("compare",{})
            mappings2 = r2.get("mappings",{})
            comp_cos_list = list(comp2.keys())
            st.markdown(f"**{b2}** 기준  |  비교: {', '.join(comp_cos_list)}")
            stat_cols = st.columns(2 + len(comp_cos_list))
            stat_cols[0].metric("기준회사 담보", f"{len(bcovs2)}개", b2)
            for i, co in enumerate(comp_cos_list):
                n_comp = len(comp2.get(co, []))
                stat_cols[i+1].metric(f"{co} 담보", f"{n_comp}개")
            matched_cnt = sum(1 for v in mappings2.get(comp_cos_list[0],{}).values() if v)
            pct = round(matched_cnt/len(bcovs2)*100) if bcovs2 else 0
            stat_cols[-1].metric("매핑 성공률", f"{matched_cnt}/{len(bcovs2)}", f"{pct}%")

            CAT_ORDER_DISP = ["암","뇌/심","수술비","입원비","치료비","통원비","기타"]
            CAT_COLORS_CSS = {
                "암":"background:#FBEAF0;color:#993556","뇌/심":"background:#E1F5EE;color:#0F6E56",
                "수술비":"background:#E6F1FB;color:#185FA5","입원비":"background:#EAF3DE;color:#3B6D11",
                "치료비":"background:#FAEEDA;color:#854F0B","통원비":"background:#EEF;color:#535",
                "기타":"background:#F1EFE8;color:#5F5E5A",
            }
            cat_cols = st.columns(2)
            for ci, co in enumerate([b2] + comp_cos_list[:1]):
                covs_for_cat = bcovs2 if co == b2 else comp2.get(co,[])
                cnt = Counter(c.get("category","기타") for c in covs_for_cat)
                badges = " ".join(
                    f'<span style="display:inline-block;padding:2px 8px;border-radius:10px;'
                    f'font-size:11px;font-weight:600;margin:2px;{CAT_COLORS_CSS.get(cat,CAT_COLORS_CSS["기타"])}">'
                    f'{cat} {cnt.get(cat,0)}개</span>'
                    for cat in CAT_ORDER_DISP if cnt.get(cat,0) > 0
                )
                cat_cols[ci].markdown(f"**{co}** 담보 분류")
                cat_cols[ci].markdown(badges, unsafe_allow_html=True)

            st.markdown("---")
            st.markdown("**담보 매핑 비교표**")
            import pandas as pd
            preview_rows = []
            for h in bcovs2:
                row = {"분류": h.get("category",""), f"{b2} 담보명": h['name']}
                for co in comp_cos_list:
                    m = mappings2.get(co,{}).get(h['name'])
                    row[f"{co} 담보명"] = m['name'] if m else "❌ 해당없음"
                preview_rows.append(row)
            df_cmp = pd.DataFrame(preview_rows)
            st.dataframe(df_cmp, use_container_width=True, height=500)

            wb_t2c = openpyxl.Workbook(); wb_t2c.remove(wb_t2c.active)
            build_sheet_comparison(wb_t2c, b2, bcovs2, comp2, mappings2)
            buf2c = io.BytesIO(); wb_t2c.save(buf2c)
            dl_col, info_col = st.columns([1, 3])
            with dl_col:
                st.download_button(
                    "⬇️ 엑셀 다운로드 (담보비교표)",
                    data=buf2c.getvalue(), file_name="담보비교표.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            with info_col:
                st.info("📋 담보별 상세 지급사유 및 면책, 감액 사항은 엑셀 파일에서 확인하세요.")


def show_tab3_content():
    companies_loaded_fn = list(st.session_state.pdf_coverages.keys())
    st.markdown("### 📋 경쟁력 강화방안")
    if not companies_loaded_fn:
        st.warning("📂 업로드된 PDF가 없습니다. 파일 업로드 메뉴에서 요약서를 먼저 업로드해주세요.")
        if st.button("📁 파일 업로드 메뉴로 이동", type="primary"):
            st.session_state.current_menu = "파일 업로드"
        return
    api_key = os.environ.get("ANTHROPIC_API_KEY","")

    avail3 = companies_loaded_fn
    col1, col2 = st.columns(2)
    with col1:
        default_idx = avail3.index("한화손해보험") if "한화손해보험" in avail3 else 0
        base_co3 = st.selectbox("기준회사", avail3, index=default_idx, key="t3_base")
    with col2:
        comp_cos3 = st.multiselect("비교회사",
            [c for c in avail3 if c != base_co3], key="t3_compare")

    if st.button("📝 보고서 생성", key="t3_run", type="primary"):
        if not api_key:
            st.error("ANTHROPIC_API_KEY 환경변수가 설정되지 않았습니다. 터미널에서 'set ANTHROPIC_API_KEY=your_key' 후 재실행하세요.")
        elif not comp_cos3:
            st.warning("비교회사를 1개 이상 선택하세요.")
        else:
            target3 = [base_co3] + comp_cos3
            struct_data = {co: st.session_state.pdf_structures.get(co,{}) for co in target3}
            cov_data    = {co: st.session_state.pdf_coverages.get(co,[]) for co in target3}

            base_names    = set(c['name'] for c in cov_data.get(base_co3,[]))
            compare_names = set()
            for co in comp_cos3:
                compare_names.update(c['name'] for c in cov_data.get(co,[]))
            base_only_covs    = [c for c in cov_data.get(base_co3,[]) if c['name'] not in compare_names][:15]
            compare_only_covs = []
            for co in comp_cos3:
                compare_only_covs += [c for c in cov_data.get(co,[]) if c['name'] not in base_names]
            compare_only_covs = compare_only_covs[:15]

            mapping_lines = []
            for co in comp_cos3:
                with st.spinner(f"🔄 {co} 매핑 중..."):
                    mapping_tmp = build_mapping(cov_data.get(base_co3,[]), cov_data.get(co,[]))
                matched_n = sum(1 for v in mapping_tmp.values() if v)
                total_n   = len(cov_data.get(base_co3,[]))
                mapping_lines.append(f"{co}: 기준회사 {total_n}개 중 {matched_n}개 매핑 ({round(matched_n/total_n*100)}%)")
            mapping_summary = "\n".join(mapping_lines)

            cat_stats = {}
            for co in target3:
                cnt = Counter(c['category'] for c in cov_data.get(co,[]))
                cat_stats[co] = dict(cnt)

            system_prompt = """당신은 손해보험사에서 영업직무 10년, 상품개발·기획부서 10년을 경험한
손해보험 상품 전략 전문가입니다.

현장 영업 경험을 바탕으로 고객 관점에서의 상품 경쟁력을 파악하고,
상품개발 전문가로서 담보 구성·보장 구조·시장 포지셔닝 관점에서
실질적인 전략 제언을 할 수 있습니다.

보고서는 아래 5개 섹션으로 구성됩니다:
1. SUMMARY: 핵심 요약 (executive_summary + key_message)
2. 상품구조 비교: 상품구조 GAP은 별도 데이터로 자동 생성 (제공 데이터 기반)
3. 담보운영 비교: coverage_gap (기준사/비교사 단독 보유 담보) + competitive_matrix (경쟁력 매트릭스)
4. SWOT 분석: swot (S강점/W약점/O기회/T위협)
5. 경쟁력 강화 액션플랜: action_plan (우선순위별 실행 계획)

작성 원칙:
1. 반드시 제공된 상품구조·담보 비교 데이터에 근거하여 작성
2. 데이터에 없는 내용은 절대 추론하거나 지어내지 말 것
3. 수치와 사실 기반으로 구체적으로 작성
4. 실무에서 즉시 활용 가능한 전략 제언 중심으로 작성
5. 응답은 반드시 아래 JSON 형식으로만 반환 (마크다운 코드블록이나 설명 텍스트 없이 순수 JSON만)
6. 각 문자열 필드는 반드시 지정된 글자 수 이내로 작성할 것
7. JSON 문자열 내부에 큰따옴표(")를 절대 사용하지 말 것. 작은따옴표(') 사용
8. action_plan 최대 3개, swot 각 항목 최대 3개, competitive_matrix 최대 8개, coverage_gap 각 항목 최대 6개"""

            base_only_json = json.dumps(
                [{"name": c['name'], "category": c['category']} for c in base_only_covs],
                ensure_ascii=False
            )
            compare_only_json = json.dumps(
                [{"name": c['name'], "category": c['category'], "reason": c.get('reason','')[:60]} for c in compare_only_covs],
                ensure_ascii=False
            )
            json_schema = """{
  "executive_summary": "3~4문장 (200자 이내)",
  "swot": {
    "strength": [{"item":"강점명(20자이내)","detail":"근거(50자이내)","impact":"영향(30자이내)"}],
    "weakness": [{"item":"약점명(20자이내)","detail":"근거(50자이내)","impact":"리스크(30자이내)"}],
    "opportunity": [{"item":"기회명(20자이내)","detail":"내용(50자이내)","action":"방향(30자이내)"}],
    "threat": [{"item":"위협명(20자이내)","detail":"열위요소(50자이내)","mitigation":"대응(30자이내)"}]
  },
  "competitive_matrix": [
    {"category":"항목명(15자이내)","base_score":4,"compare_scores":{"비교회사명":3},"comment":"근거(40자이내)"}
  ],
  "coverage_gap": {
    "base_only": [{"name":"담보명(30자이내)","category":"분류","strategic_value":"가치(30자이내)"}],
    "compare_only": [{"name":"담보명(30자이내)","category":"분류","adoption_priority":"high/medium/low","reason":"이유(30자이내)"}]
  },
  "action_plan": [
    {"priority":1,"title":"제목(20자이내)","category":"담보신설/구조개선/서비스강화/가격전략",
      "what":"내용(60자이내)","why":"근거(60자이내)","expected_effect":"효과(40자이내)",
      "timeline":"단기(3개월이내)/중기(6개월이내)/장기(1년이내)"}
  ],
  "key_message": "30자이내"
}"""

            user_prompt = f"""아래 데이터를 바탕으로 [{base_co3}] 경쟁력 분석 보고서를 작성하세요.
비교 대상: {', '.join(comp_cos3)}

## 탭1 상품구조 데이터
{json.dumps({co: {k:v for k,v in d.items() if k!='company'} for co,d in struct_data.items()}, ensure_ascii=False, indent=2)[:2500]}

## 탭2 담보 비교 데이터
담보 수: {json.dumps({co: len(cov_data.get(co,[])) for co in target3}, ensure_ascii=False)}
분류별 담보 수: {json.dumps(cat_stats, ensure_ascii=False)}
매핑 결과:
{mapping_summary}

기준회사만 보유한 담보 (상위 15개):
{base_only_json}

비교회사만 보유한 담보 (상위 15개):
{compare_only_json}

아래 JSON 구조로 정확히 반환하세요 (순수 JSON만, 다른 텍스트 없이):
{json_schema}"""

            with st.spinner("🤖 손해보험 전문가 AI가 보고서를 작성 중입니다..."):
                try:
                    try:
                        import anthropic as _anthropic_mod
                    except ImportError:
                        st.error(
                            "❌ anthropic 패키지가 설치되지 않았습니다.\n\n"
                            "터미널에서 아래 명령어를 실행 후 Streamlit을 재시작하세요:\n\n"
                            "```\npip install anthropic\n```"
                        )
                        st.stop()
                    import httpx as _httpx
                    import warnings as _warnings
                    _warnings.filterwarnings("ignore", message="Unverified HTTPS request")
                    client = _anthropic_mod.Anthropic(
                        api_key=api_key,
                        http_client=_httpx.Client(
                            timeout=_httpx.Timeout(120.0, connect=10.0),
                            verify=False  # 사내 SSL 인터셉트 프록시 우회
                        )
                    )
                    resp = client.messages.create(
                        model="claude-sonnet-4-5",
                        max_tokens=8000,
                        system=system_prompt,
                        messages=[{"role":"user","content":user_prompt}]
                    )
                    raw_text = resp.content[0].text.strip()
                    raw_text = re.sub(r'^```json\s*', '', raw_text, flags=re.MULTILINE)
                    raw_text = re.sub(r'^```\s*', '', raw_text, flags=re.MULTILINE)
                    raw_text = re.sub(r'\s*```\s*$', '', raw_text.strip())
                    raw_text = raw_text.strip()
                    json_start = raw_text.find('{')
                    json_end   = raw_text.rfind('}')
                    if json_start >= 0 and json_end > json_start:
                        raw_text = raw_text[json_start:json_end+1]
                    report_json = json.loads(raw_text)
                    st.session_state.tab3_report = {
                        "json": report_json,
                        "base": base_co3,
                        "compare": comp_cos3,
                        "struct_data": struct_data,
                    }
                except json.JSONDecodeError as e:
                    st.error(f"보고서 생성 실패 - JSON 파싱 오류: {e}")
                    st.text_area("원본 응답 (디버그)", raw_text[:1000])
                except Exception as e:
                    err_msg = str(e)
                    err_type = type(e).__name__
                    if "AuthenticationError" in err_type or "401" in err_msg or "invalid_api_key" in err_msg:
                        st.error(
                            "❌ API 키 오류\n\n"
                            "ANTHROPIC_API_KEY 환경변수 값을 확인하세요.\n\n"
                            f"상세 오류: {err_msg}"
                        )
                    elif "Connection" in err_msg or "connect" in err_msg.lower() or "ConnectError" in err_type:
                        st.error(
                            "❌ API 연결 오류\n\n"
                            "해결 방법:\n"
                            "1. 인터넷 연결 확인\n"
                            "2. 방화벽/프록시 설정 확인\n"
                            "3. `pip install -U anthropic httpx` 실행 후 재시작\n\n"
                            f"상세 오류: {err_msg}"
                        )
                    elif "RateLimitError" in err_type or "rate_limit" in err_msg:
                        st.error(f"❌ API 요청 한도 초과. 잠시 후 다시 시도하세요.\n\n상세 오류: {err_msg}")
                    else:
                        st.error(f"❌ API 오류 ({err_type}): {err_msg}")
                    import traceback
                    with st.expander("🔧 상세 오류 내용 (디버그)"):
                        st.code(traceback.format_exc())

    if st.session_state.get("tab3_report"):
        r3 = st.session_state.tab3_report
        report_json = r3.get("json",{})
        base_r  = r3.get("base","")
        comp_r  = r3.get("compare",[])
        struct_r = r3.get("struct_data", {})
        html_report = render_report_html(report_json, base_r, comp_r, struct_r)
        st.components.v1.html(html_report, height=3200, scrolling=True)
        word_bytes = save_word_from_json(report_json, base_r, comp_r)
        if word_bytes:
            st.download_button(
                "⬇️ Word 다운로드",
                data=word_bytes,
                file_name=f"경쟁력강화방안_{base_r}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )


# ════════════════════════════════════════════════════════════
#  H. render_report_html 함수
# ════════════════════════════════════════════════════════════
def render_report_html(report_data, base_co, comp_cos, struct_data=None):
    """JSON 보고서 데이터를 A4 보고서 형태 HTML로 변환"""
    swot    = report_data.get("swot", {})
    matrix  = report_data.get("competitive_matrix", [])
    gap     = report_data.get("coverage_gap", {})
    actions = report_data.get("action_plan", [])
    comp_label = ", ".join(comp_cos)

    # ── 헬퍼: SWOT 항목 ──────────────────────────────────────
    def swot_items_html(items):
        if not items:
            return '<p class="swot-item" style="color:#999">데이터 미확인</p>'
        out = ""
        for item in items:
            detail = item.get("detail","")
            impact = item.get("impact") or item.get("action") or item.get("mitigation","")
            out += f"""<div class="swot-item">
              <span class="swot-item-title">· {item.get('item','')}</span>
              <span class="swot-item-detail"> — {detail}</span>
              {f'<div style="font-size:11px;color:#888;margin-left:12px;font-style:italic">{impact}</div>' if impact else ''}
            </div>"""
        return out

    # ── 헬퍼: 별점 ───────────────────────────────────────────
    def star(score, hi=True):
        s = max(0, min(5, int(score)))
        col = "#FF6B00" if hi else "#C62828"
        return (f'<span style="color:{col}">{"★"*s}</span>'
                f'<span style="color:#ddd">{"☆"*(5-s)}</span>')

    # ── 헬퍼: 분류 뱃지 색상 ─────────────────────────────────
    CAT_BADGE = {
        "암":    ("FBEAF0","993556"), "뇌/심": ("E1F5EE","0F6E56"),
        "수술비":("E6F1FB","185FA5"), "입원비": ("EAF3DE","3B6D11"),
        "치료비":("FAEEDA","854F0B"), "통원비": ("EEF","535"),
        "기타":  ("F1EFE8","5F5E5A"),
    }
    def cat_badge(cat):
        bg, fg = CAT_BADGE.get(cat, ("F1EFE8","5F5E5A"))
        return f'<span style="background:#{bg};color:#{fg};padding:1px 7px;border-radius:10px;font-size:10px;font-weight:600">{cat}</span>'

    # ── 경쟁력 매트릭스 행 ───────────────────────────────────
    matrix_rows = ""
    for row in matrix:
        base_sc = row.get("base_score", 3)
        comp_sc_val = list(row.get("compare_scores", {}).values())[0] if row.get("compare_scores") else 3
        hi = base_sc >= comp_sc_val
        bg_row = "#F1FAF4" if hi else "#FEF3F3"
        matrix_rows += f"""<tr style="background:{bg_row}">
          <td style="font-size:13px;padding:8px 12px;font-weight:500">{row.get('category','')}</td>
          <td style="padding:8px 12px;text-align:center;font-size:15px">{star(base_sc, hi)}</td>
          <td style="padding:8px 12px;text-align:center;font-size:15px">{star(comp_sc_val, not hi)}</td>
          <td style="font-size:12px;color:#555;padding:8px 12px">{row.get('comment','')}</td>
        </tr>"""

    # ── 상품구조 GAP 분석 ────────────────────────────────────
    STRUCT_LABELS = {
        "product_name":"상품명", "jong":"종(種)", "hyung":"형(型)",
        "insurance_period":"보험기간", "payment_period":"납입기간",
        "age_range":"연령", "premium_waiver":"납입면제",
        "no_claim":"무사고전환", "services":"부가서비스", "discounts":"할인",
    }
    struct_gap_rows = ""
    if struct_data and len(struct_data) >= 2:
        all_cos = list(struct_data.keys())
        base_s  = struct_data.get(base_co, {})
        # 비교사가 여럿일 경우 첫 번째만 상세 비교, 나머지는 참고
        comp_s_list = [(co, struct_data.get(co,{})) for co in all_cos if co != base_co]
        for key, label in STRUCT_LABELS.items():
            b_val = base_s.get(key,"미기재") or "미기재"
            comp_cells = ""
            diff_flag = False
            for co, cs in comp_s_list:
                c_val = cs.get(key,"미기재") or "미기재"
                is_diff = (b_val.strip() != c_val.strip()
                           and b_val != "미기재" and c_val != "미기재")
                if is_diff:
                    diff_flag = True
                cell_style = "background:#FEF3F3;" if is_diff else ""
                comp_cells += f'<td style="font-size:12px;padding:7px 10px;{cell_style}">{c_val}</td>'
            row_bg = "#FFFDE7" if diff_flag else "#FFFFFF"
            diff_marker = '<span style="color:#E65100;font-size:10px;font-weight:700;margin-left:4px">◀ 차이</span>' if diff_flag else ''
            struct_gap_rows += f"""<tr style="background:{row_bg}">
              <td style="font-size:12px;padding:7px 10px;font-weight:600;color:#FF6B00;
                  background:#FFF8F3;border-right:1px solid #eee">{label}</td>
              <td style="font-size:12px;padding:7px 10px">{b_val}{diff_marker}</td>
              {comp_cells}
            </tr>"""
        # 헤더
        comp_headers = "".join(
            f'<th style="font-size:12px;padding:8px 10px;background:#F5F5F5;'
            f'border-bottom:2px solid #ddd;color:#555">{co}</th>'
            for co,_ in comp_s_list
        )
        struct_gap_section = f"""
<table class="matrix-table">
  <thead><tr>
    <th style="width:110px;background:#FF6B00;color:#fff;border-bottom:2px solid #FF6B00">항목</th>
    <th style="font-size:12px;padding:8px 10px;background:#FFF3E8;
        border-bottom:2px solid #FF6B00;color:#FF6B00;font-weight:700">{base_co}</th>
    {comp_headers}
  </tr></thead>
  <tbody>{struct_gap_rows}</tbody>
</table>
<div style="font-size:11px;color:#888;margin-top:6px">
  ※ 노란 행: 기준회사와 비교회사 간 차이가 있는 항목
</div>"""
    else:
        struct_gap_section = ""

    struct_gap_section_inner = struct_gap_section  # 보고서 섹션2에 사용

    # ── 담보 GAP 항목 ────────────────────────────────────────
    def gap_item_html(item, side):
        if side == "base":
            cat  = item.get("category","")
            val  = item.get("strategic_value","")
            name = item.get("name","")
            return f"""<div class="gap-item">
              {cat_badge(cat)}
              <div><div style="font-size:12px;font-weight:600">{name}</div>
              <div style="font-size:11px;color:#666;margin-top:1px">{val}</div></div>
            </div>"""
        else:
            pri  = item.get("adoption_priority","low")
            name = item.get("name","")
            rsn  = item.get("reason","")
            pri_colors = {"high":("FDEAEE","C62828"),"medium":("FFFAE0","E65100"),"low":("F1EFE8","5F5E5A")}
            bg, fg = pri_colors.get(pri, ("F1EFE8","5F5E5A"))
            return f"""<div class="gap-item">
              <span style="background:#{bg};color:#{fg};padding:2px 7px;border-radius:10px;
                  font-size:10px;font-weight:700;white-space:nowrap;flex-shrink:0">{pri.upper()}</span>
              <div><div style="font-size:12px;font-weight:600">{name}</div>
              <div style="font-size:11px;color:#666;margin-top:1px">{rsn}</div></div>
            </div>"""

    base_gap_html    = "".join(gap_item_html(i,"base")    for i in gap.get("base_only",[]))    or '<p style="font-size:12px;color:#999">해당 없음</p>'
    compare_gap_html = "".join(gap_item_html(i,"compare") for i in gap.get("compare_only",[])) or '<p style="font-size:12px;color:#999">해당 없음</p>'

    # ── Action Plan 항목 ─────────────────────────────────────
    tl_colors  = {"단기(3개월이내)":"#C62828","중기(6개월이내)":"#E65100","장기(1년이내)":"#757575"}
    cat_colors = {"담보신설":"#1565C0","구조개선":"#2E7D32","서비스강화":"#6A1B9A","가격전략":"#E65100"}
    action_items_html = ""
    for a in sorted(actions, key=lambda x: x.get("priority",9)):
        tc = tl_colors.get(a.get("timeline",""), "#757575")
        cc = cat_colors.get(a.get("category",""), "#888")
        action_items_html += f"""
        <div class="action-item">
          <div class="action-header">
            <div class="action-num">P{a.get('priority','')}</div>
            <span class="action-title">{a.get('title','')}</span>
            <span style="background:{tc};color:#fff;padding:2px 9px;border-radius:10px;font-size:11px;font-weight:600">{a.get('timeline','')}</span>
            <span style="background:{cc};color:#fff;padding:2px 9px;border-radius:10px;font-size:11px;font-weight:600">{a.get('category','')}</span>
          </div>
          <div class="action-detail">
            <div class="action-detail-row"><span style="color:#FF6B00;font-weight:700">▸ WHY</span>&nbsp; {a.get('why','')}</div>
            <div class="action-detail-row"><span style="color:#1565C0;font-weight:700">▸ WHAT</span>&nbsp; {a.get('what','')}</div>
            <div class="action-detail-row"><span style="color:#2E7D32;font-weight:700">▸ 기대효과</span>&nbsp; {a.get('expected_effect','')}</div>
          </div>
        </div>"""

    # ── 최종 HTML ────────────────────────────────────────────
    html = f"""<!DOCTYPE html><html><head>
<meta charset="utf-8">
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{
  font-family:Arial,'Malgun Gothic','맑은 고딕',sans-serif;
  background:#EFEFEF;padding:28px 0;color:#222;line-height:1.6;
}}
.report{{
  max-width:820px;margin:0 auto;background:#fff;
  padding:48px 56px;
  box-shadow:0 4px 24px rgba(0,0,0,.10);
}}
/* 보고서 상단 헤더 */
.report-header{{
  border-left:8px solid #FF6B00;
  padding:0 0 16px 20px;
  margin-bottom:24px;
}}
.report-title{{font-size:24px;font-weight:700;color:#1A1A2E;margin-bottom:6px}}
.report-meta{{font-size:13px;color:#777;margin-bottom:12px}}
.report-key-msg{{font-size:16px;font-weight:700;color:#FF6B00;margin-top:8px}}
.report-key-divider{{border:none;border-top:2px solid #FF6B00;margin-top:10px}}
/* 섹션 */
.section-title{{
  font-size:17px;font-weight:700;color:#222;
  border-left:4px solid #FF6B00;padding-left:12px;
  margin:36px 0 16px;
}}
.divider{{border:none;border-top:1px solid #EEEEEE;margin:28px 0}}
/* Executive Summary */
.exec-summary{{
  background:#FFF8F3;border-left:4px solid #FF6B00;
  padding:16px 20px;border-radius:0 6px 6px 0;
  font-size:14px;line-height:1.8;color:#222;
}}
/* SWOT */
.swot-grid{{
  display:grid;grid-template-columns:1fr 1fr;
  border:1px solid #E0E0E0;border-radius:8px;overflow:hidden;
}}
.swot-cell{{padding:16px 18px}}
.swot-header{{font-weight:700;font-size:13px;margin-bottom:10px}}
.swot-item{{font-size:13px;margin-bottom:7px;line-height:1.5}}
.swot-item-title{{font-weight:600}}
.swot-item-detail{{color:#555}}
/* 매트릭스 */
.matrix-table{{width:100%;border-collapse:collapse;font-size:13px}}
.matrix-table th{{
  background:#F5F5F5;padding:9px 12px;text-align:left;
  border-bottom:2px solid #FF6B00;font-weight:600;color:#333;
}}
.matrix-table td{{padding:8px 12px;border-bottom:1px solid #EEEEEE;vertical-align:top}}
/* GAP */
.gap-grid{{display:grid;grid-template-columns:1fr 1fr;gap:20px}}
.gap-section-title{{font-size:13px;font-weight:700;margin-bottom:10px;padding-bottom:6px;border-bottom:1px solid #eee}}
.gap-item{{font-size:12px;margin-bottom:10px;display:flex;gap:8px;align-items:flex-start}}
/* Action Plan */
.action-item{{margin-bottom:22px;padding-bottom:18px;border-bottom:1px solid #F0F0F0}}
.action-item:last-child{{border-bottom:none;margin-bottom:0}}
.action-header{{display:flex;align-items:center;gap:10px;margin-bottom:10px;flex-wrap:wrap}}
.action-num{{
  background:#FF6B00;color:#fff;border-radius:50%;
  width:26px;height:26px;display:flex;align-items:center;
  justify-content:center;font-weight:700;font-size:13px;flex-shrink:0;
}}
.action-title{{font-weight:700;font-size:15px;color:#111}}
.action-detail{{font-size:13px;color:#444;margin-left:36px;line-height:1.8}}
.action-detail-row{{margin-bottom:4px}}
/* 푸터 */
.report-footer{{
  margin-top:40px;padding-top:16px;
  border-top:1px solid #ddd;
  font-size:11px;color:#999;text-align:right;line-height:1.8;
}}
</style>
</head><body>
<div class="report">

  <!-- 보고서 헤더 -->
  <div class="report-header">
    <div class="report-title">한화손해보험 경쟁력 강화방안 보고서</div>
    <div class="report-meta">기준회사: <strong>{base_co}</strong>&nbsp;&nbsp;|&nbsp;&nbsp;비교대상: <strong>{comp_label}</strong></div>
    <hr class="report-key-divider">
    <div class="report-key-msg">🏆 {report_data.get('key_message','')}</div>
  </div>

  <!-- SECTION 1: SUMMARY -->
  <div class="section-title">1. SUMMARY</div>
  <div class="exec-summary">{report_data.get('executive_summary','')}</div>

  <!-- SECTION 2: 상품구조 비교 -->
  <hr class="divider">
  <div class="section-title">2. 상품구조 비교</div>
  {struct_gap_section_inner}

  <!-- SECTION 3: 담보운영 비교 -->
  <hr class="divider">
  <div class="section-title">3. 담보운영 비교</div>
  <div class="gap-grid">
    <div>
      <div class="gap-section-title" style="color:#2E7D32">✅ {base_co} 단독 보유 — 전략적 강점 담보</div>
      {base_gap_html}
    </div>
    <div>
      <div class="gap-section-title" style="color:#C62828">🔺 비교사 단독 보유 — 도입 검토 담보</div>
      {compare_gap_html}
    </div>
  </div>
  <hr class="divider">
  <div class="section-title" style="margin-top:0">경쟁력 매트릭스 (★ 5점 만점)</div>
  <table class="matrix-table">
    <thead><tr>
      <th>항목</th>
      <th style="text-align:center;color:#FF6B00">{base_co}</th>
      <th style="text-align:center;color:#666">{comp_label}</th>
      <th>평가 근거</th>
    </tr></thead>
    <tbody>{matrix_rows}</tbody>
  </table>

  <!-- SECTION 4: SWOT 분석 -->
  <hr class="divider">
  <div class="section-title">4. SWOT 분석 — {base_co}</div>
  <div class="swot-grid">
    <div class="swot-cell" style="background:#D4F0D4;border-right:1px solid #C8E6C9;border-bottom:1px solid #C8E6C9">
      <div class="swot-header" style="color:#2E7D32">S &nbsp;강점</div>
      {swot_items_html(swot.get('strength',[]))}
    </div>
    <div class="swot-cell" style="background:#FDEAEE;border-bottom:1px solid #FFCDD2">
      <div class="swot-header" style="color:#C62828">W &nbsp;약점</div>
      {swot_items_html(swot.get('weakness',[]))}
    </div>
    <div class="swot-cell" style="background:#EAF4FD;border-right:1px solid #BBDEFB">
      <div class="swot-header" style="color:#1565C0">O &nbsp;기회</div>
      {swot_items_html(swot.get('opportunity',[]))}
    </div>
    <div class="swot-cell" style="background:#FFFAE0">
      <div class="swot-header" style="color:#E65100">T &nbsp;위협</div>
      {swot_items_html(swot.get('threat',[]))}
    </div>
  </div>

  <!-- SECTION 5: 경쟁력 강화 Action Plan -->
  <hr class="divider">
  <div class="section-title">5. 경쟁력 강화 액션플랜</div>
  {action_items_html}

  <!-- 푸터 -->
  <div class="report-footer">
    본 보고서는 업로드된 PDF 요약서 원문 데이터에 근거하여 작성되었습니다.<br>
    손해보험 영업·상품개발 전문가 관점 분석 &nbsp;|&nbsp; Claude AI<br>
    <strong>HANWHA Insurance &nbsp;|&nbsp; 상품개발부</strong>
  </div>

</div>
</body></html>"""
    return html


def save_word_from_json(report_data, base_co, comp_cos):
    """JSON 보고서를 Word 문서로 저장"""
    try:
        from docx import Document
        from docx.shared import RGBColor, Pt
        doc = Document()
        title = doc.add_heading(f"경쟁력 강화 방안 보고서 — {base_co}", 0)
        if title.runs:
            title.runs[0].font.color.rgb = RGBColor(0xFF,0x6B,0x00)
        doc.add_paragraph(f"비교 대상: {', '.join(comp_cos)}")
        doc.add_paragraph(f"핵심 메시지: {report_data.get('key_message','')}")
        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph(report_data.get("executive_summary",""))
        swot = report_data.get("swot",{})
        doc.add_heading("SWOT 분석", 1)
        for key, label in [("strength","S 강점"),("weakness","W 약점"),("opportunity","O 기회"),("threat","T 위협")]:
            doc.add_heading(label, 2)
            for item in swot.get(key,[]):
                doc.add_paragraph(f"{item.get('item','')}: {item.get('detail','')}", style="List Bullet")
        doc.add_heading("경쟁력 매트릭스", 1)
        for row in report_data.get("competitive_matrix",[]):
            doc.add_paragraph(f"{row.get('category','')}: {base_co} {row.get('base_score','')}점 / {row.get('comment','')}", style="List Bullet")
        doc.add_heading("담보 GAP 분석", 1)
        doc.add_heading(f"{base_co} 단독 보유", 2)
        for item in report_data.get("coverage_gap",{}).get("base_only",[]):
            doc.add_paragraph(f"{item.get('name','')}: {item.get('strategic_value','')}", style="List Bullet")
        doc.add_heading("비교사 단독 보유 (도입 검토)", 2)
        for item in report_data.get("coverage_gap",{}).get("compare_only",[]):
            doc.add_paragraph(f"[{item.get('adoption_priority','').upper()}] {item.get('name','')}: {item.get('reason','')}", style="List Bullet")
        doc.add_heading("Action Plan", 1)
        for a in sorted(report_data.get("action_plan",[]), key=lambda x: x.get("priority",9)):
            doc.add_heading(f"P{a.get('priority','')}. {a.get('title','')}", 2)
            doc.add_paragraph(f"타임라인: {a.get('timeline','')}")
            doc.add_paragraph(f"WHY: {a.get('why','')}")
            doc.add_paragraph(f"WHAT: {a.get('what','')}")
            doc.add_paragraph(f"기대효과: {a.get('expected_effect','')}")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        return None


# ════════════════════════════════════════════════════════════
#  G. 메인 라우터
# ════════════════════════════════════════════════════════════
menu = st.session_state.current_menu
if menu == "파일 업로드":
    show_upload_page()
elif menu == "상품 구조 비교":
    show_tab1_content()
elif menu == "담보 내용 비교":
    show_tab2_content()
elif menu == "경쟁력 강화방안":
    show_tab3_content()
